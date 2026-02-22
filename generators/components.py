"""
Reusable Word document building blocks for worksheet generation.
Extracted and generalised from create_myth_worksheet.py.
All components are theme-aware and support differentiation levels.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from generators.styles import FONT_NAME, COLOURS, WORD_TYPES, THEMES, DIFF_LEVELS


# ─── Low-Level Helpers ─────────────────────────────────────────────────────────


def set_run_font(run, font_name=FONT_NAME, size=Pt(14), bold=False, italic=False, colour=None):
    """Apply consistent font formatting to a run."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>'
    )
    existing = rPr.findall(f'{{{rPr.nsmap["w"]}}}rFonts')
    for e in existing:
        rPr.remove(e)
    rPr.insert(0, rFonts)


def set_cell_shading(cell, colour_hex):
    """Set background colour of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{colour_hex}" w:val="clear"/>'
    )
    existing = tcPr.findall(f'{{{tcPr.nsmap["w"]}}}shd')
    for e in existing:
        tcPr.remove(e)
    tcPr.append(shading)


def set_cell_borders(cell, colour_hex, sz=8):
    """Set all borders of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'</w:tcBorders>'
    )
    existing = tcPr.findall(f'{{{tcPr.nsmap["w"]}}}tcBorders')
    for e in existing:
        tcPr.remove(e)
    tcPr.append(borders)


def set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    """Set cell padding in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing = tcPr.findall(f'{{{tcPr.nsmap["w"]}}}tcMar')
    for e in existing:
        tcPr.remove(e)
    tcPr.append(margins)


def set_table_full_width(table):
    """Make table span full page width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    existing = tblPr.findall(f'{{{tblPr.nsmap["w"]}}}tblW')
    for e in existing:
        tblPr.remove(e)
    tblPr.append(tblW)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.findall(f'{{{tblPr.nsmap["w"]}}}tblBorders')
    for e in existing:
        tblPr.remove(e)
    tblPr.append(borders)


def set_no_spacing(paragraph):
    """Remove extra spacing above/below a paragraph."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


# ─── Document Setup ────────────────────────────────────────────────────────────


def create_base_document(extra_spacing=False):
    """Create a new Document with standard page margins and default font."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(14)
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
    )
    existing = rPr.findall(f'{{{rPr.nsmap["w"]}}}rFonts')
    for e in existing:
        rPr.remove(e)
    rPr.insert(0, rFonts)

    if extra_spacing:
        style.paragraph_format.line_spacing = Pt(36)

    # Remove the default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    return doc


# ─── High-Level Components ─────────────────────────────────────────────────────


def add_title_area(doc, title, subtitle, theme_key='classic', level='expected', is_answer_key=False):
    """Add a themed title block with name/date fields."""
    theme = THEMES[theme_key]
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'

    if is_answer_key:
        title = f"{title} - ANSWER KEY"
        subtitle = "Teacher Edition"

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, theme['body'])
    set_cell_borders(cell, theme['header'], sz=12)
    set_cell_padding(cell, top=150, bottom=150, left=200, right=200)

    title_size = Pt(28) if is_dev else Pt(24)

    # Title with theme icon
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p)
    p.paragraph_format.space_after = Pt(4)
    full_title = f"{theme['icon']} {title}"
    run = p.add_run(full_title)
    set_run_font(run, size=title_size, bold=True, colour=COLOURS['title_text'])

    # Subtitle
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p2)
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(subtitle)
    set_run_font(run2, size=Pt(14) if is_dev else Pt(12), italic=True, colour=COLOURS['grey_text'])

    # Name and Date line (skip for answer keys)
    if not is_answer_key:
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_no_spacing(p3)
        run3 = p3.add_run('Name: ________________________          Date: ________________')
        set_run_font(run3, size=Pt(16) if is_dev else Pt(14), colour=COLOURS['black'])


def add_learning_objective(doc, objective, theme_key='classic'):
    """Add a learning objective box."""
    theme = THEMES[theme_key]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F5F5')
    set_cell_borders(cell, theme['accent'], sz=8)
    set_cell_padding(cell, top=80, bottom=80, left=150, right=150)

    p = cell.paragraphs[0]
    set_no_spacing(p)
    r1 = p.add_run('Learning Objective: ')
    set_run_font(r1, size=Pt(11), bold=True, colour=COLOURS['grey_text'])
    r2 = p.add_run(objective)
    set_run_font(r2, size=Pt(11), italic=True, colour=COLOURS['grey_text'])


def add_colour_key(doc, level='expected', word_types_to_show=None):
    """Add a colour/symbol key legend."""
    if word_types_to_show is None:
        word_types_to_show = ['time', 'adjective', 'verb', 'noun']
        if level != 'developing':
            word_types_to_show.extend(['name', 'open'])

    # Filter out unknown word types to prevent KeyError
    word_types_to_show = [wt for wt in word_types_to_show if wt in WORD_TYPES]
    if not word_types_to_show:
        return

    is_dev = level == 'developing'
    key_table = doc.add_table(rows=1, cols=len(word_types_to_show))
    set_table_full_width(key_table)
    remove_table_borders(key_table)

    for i, wt_key in enumerate(word_types_to_show):
        wt = WORD_TYPES[wt_key]
        cell = key_table.cell(0, i)
        set_cell_shading(cell, wt['bg'])
        set_cell_borders(cell, wt['border'], sz=6)
        set_cell_padding(cell, top=40, bottom=40, left=40, right=40)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        font_size = Pt(11) if is_dev else Pt(9)
        run = p.add_run(f'{wt["symbol"]} {wt["label"]}')
        set_run_font(run, size=font_size, bold=True, colour=wt['text'])


def add_instructions(doc, text, level='expected'):
    """Add instruction text with colour key."""
    diff = DIFF_LEVELS[level]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    set_no_spacing(p)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=Pt(diff['font_size'] - 2), italic=True, colour=COLOURS['grey_text'])

    add_colour_key(doc, level)


def add_section_header(doc, section_number, title, theme_key='classic'):
    """Add a themed, coloured section header bar."""
    theme = THEMES[theme_key]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(10)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, theme['header'])
    set_cell_borders(cell, theme['header'], sz=4)
    set_cell_padding(cell, top=80, bottom=80, left=200, right=200)

    full_title = f"{theme['icon']} {theme['section']} {section_number}: {title}"

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p)
    run = p.add_run(full_title)
    set_run_font(run, size=Pt(18), bold=True, colour=COLOURS['white'])


def add_reminder_box(doc, text, theme_key='classic'):
    """Add a themed reminder prompt box."""
    theme = THEMES[theme_key]

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLOURS['reminder_bg'])
    set_cell_borders(cell, COLOURS['reminder_border'], sz=6)
    set_cell_padding(cell, top=80, bottom=80, left=150, right=150)

    p = cell.paragraphs[0]
    set_no_spacing(p)
    r1 = p.add_run(f"{theme['reminder']}! ")
    set_run_font(r1, size=Pt(12), bold=True, colour=COLOURS['reminder_text'])
    r2 = p.add_run(text)
    set_run_font(r2, size=Pt(12), italic=True, colour=COLOURS['reminder_text'])


def add_word_bank(doc, word_bank_data, level='expected'):
    """
    Add a categorised word bank with colour-coded boxes.

    word_bank_data: list of dicts with keys:
        - word_type: str (key into WORD_TYPES)
        - label: str (category header text)
        - words: list of str OR list of {"word": str, "definition": str}
    """
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'
    with_definitions = level == 'developing'

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p_title)
    p_title.paragraph_format.space_after = Pt(6)
    run = p_title.add_run('Word Bank')
    set_run_font(run, size=Pt(18 if is_dev else 16), bold=True, colour=COLOURS['black'])

    # Key line
    p_key = doc.add_paragraph()
    p_key.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p_key)
    p_key.paragraph_format.space_after = Pt(6)
    run_key = p_key.add_run('Match the colour and symbol to find the right word!')
    set_run_font(run_key, size=Pt(11), italic=True, colour=COLOURS['grey_text'])

    # Grid layout
    num_cats = len(word_bank_data)
    num_rows = (num_cats + 1) // 2
    grid = doc.add_table(rows=num_rows, cols=2)
    set_table_full_width(grid)
    remove_table_borders(grid)

    tblPr = grid._tbl.tblPr
    cell_spacing = parse_xml(f'<w:tblCellSpacing {nsdecls("w")} w:w="40" w:type="dxa"/>')
    tblPr.append(cell_spacing)

    font_size = diff['font_size'] - 2

    for i, category in enumerate(word_bank_data):
        row_idx = i // 2
        col_idx = i % 2
        cell = grid.cell(row_idx, col_idx)

        wt_key = category['word_type']
        wt = WORD_TYPES.get(wt_key, WORD_TYPES['noun'])
        set_cell_shading(cell, wt['bg'])
        set_cell_borders(cell, wt['border'], sz=10)
        set_cell_padding(cell, top=80, bottom=80, left=120, right=120)

        # Category header
        p = cell.paragraphs[0]
        set_no_spacing(p)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(category['label'])
        set_run_font(run, size=Pt(font_size), bold=True, colour=wt['text'])

        words = category['words']
        if words and isinstance(words[0], dict) and 'definition' in words[0]:
            # Words with definitions
            for item in words:
                pw = cell.add_paragraph()
                set_no_spacing(pw)
                pw.paragraph_format.space_before = Pt(1)
                pw.paragraph_format.space_after = Pt(1)
                r1 = pw.add_run(f'  {item["word"]}')
                set_run_font(r1, size=Pt(font_size - 1), bold=True, colour=COLOURS['black'])
                r2 = pw.add_run(f' \u2014 {item["definition"]}')
                set_run_font(r2, size=Pt(max(font_size - 3, 9)), italic=True, colour=COLOURS['hint_text'])
        elif words and isinstance(words[0], dict):
            # Words without definitions
            word_list = [item['word'] for item in words]
            pw = cell.add_paragraph()
            set_no_spacing(pw)
            r = pw.add_run('  ' + '  |  '.join(word_list))
            set_run_font(r, size=Pt(font_size - 1), colour=COLOURS['black'])
        elif words and isinstance(words[0], str):
            # Plain string list
            pw = cell.add_paragraph()
            set_no_spacing(pw)
            r = pw.add_run('  ' + '  |  '.join(words))
            set_run_font(r, size=Pt(font_size - 1), colour=COLOURS['black'])

    # Clear empty cells if odd number of categories
    if num_cats % 2 == 1:
        empty_cell = grid.cell(num_rows - 1, 1)
        empty_cell.paragraphs[0].clear()


def add_cloze_paragraph(cell_or_doc, pieces, level='expected', is_first=False, show_answers=False):
    """
    Add a single cloze paragraph with inline blanks.

    pieces: list of dicts:
        - {"type": "text", "text": str}
        - {"type": "blank", "word_type": str, "hint": str}           (expected/greater_depth)
        - {"type": "blank", "word_type": str, "choices": [str, ...]}  (developing)
    """
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'

    if is_first and hasattr(cell_or_doc, 'paragraphs'):
        p = cell_or_doc.paragraphs[0]
    else:
        p = cell_or_doc.add_paragraph()

    body_size = Pt(diff['font_size'])
    line_height = Pt(diff['line_spacing'])

    p.paragraph_format.space_before = Pt(6) if is_dev else Pt(4)
    p.paragraph_format.space_after = Pt(4) if is_dev else Pt(2)
    p.paragraph_format.line_spacing = line_height

    for piece in pieces:
        if piece['type'] == 'text':
            if piece.get('text'):
                run = p.add_run(piece['text'])
                set_run_font(run, size=body_size, colour=COLOURS['black'])

        elif piece['type'] == 'blank':
            word_type = piece.get('word_type', 'open')
            wt = WORD_TYPES.get(word_type, WORD_TYPES['open'])
            choices = piece.get('choices')
            hint = piece.get('hint', '')

            if show_answers:
                # Answer key mode: show the answer word instead of a blank
                answer_text = piece.get('answer', '[answer not provided]')
                answer_run = p.add_run(f' [{answer_text}] ')
                set_run_font(answer_run, size=body_size, bold=True, colour=wt['text'])
            else:
                # Student mode: show blank with hints/choices
                # Symbol prefix
                symbol_run = p.add_run(f' {wt["symbol"]}')
                set_run_font(symbol_run, size=body_size, bold=True, colour=wt['text'])

                # Underline blank
                blank_run = p.add_run(' ________________ ')
                set_run_font(blank_run, size=body_size, bold=True, colour=wt['text'])

                if is_dev and choices:
                    # Developing: show word choices
                    br_run = p.add_run()
                    br_run.add_break()
                    choices_text = '     ' + wt['symbol'] + ' Choose:  ' + '   /   '.join(choices)
                    choice_run = p.add_run(choices_text)
                    set_run_font(choice_run, size=Pt(14), bold=True, italic=True, colour=wt['text'])
                    br_run2 = p.add_run()
                    br_run2.add_break()
                elif hint:
                    # Expected/Greater Depth: show hint
                    br_run = p.add_run()
                    br_run.add_break()
                    hint_with_label = f'     {wt["symbol"]} {wt["label"]}: {hint}'
                    hint_run = p.add_run(hint_with_label)
                    hint_size = Pt(10) if level == 'greater_depth' else Pt(9)
                    set_run_font(hint_run, size=hint_size, italic=True, colour=wt['text'])
                    br_run2 = p.add_run()
                    br_run2.add_break()


def add_section_body(doc, paragraphs_data, theme_key='classic', level='expected', show_answers=False):
    """Add cloze passage paragraphs inside a themed, shaded container."""
    theme = THEMES[theme_key]
    diff = DIFF_LEVELS[level]

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, theme['body'])
    padding = diff['padding']
    set_cell_padding(cell, top=padding, bottom=padding, left=200, right=200)

    for i, para_pieces in enumerate(paragraphs_data):
        add_cloze_paragraph(cell, para_pieces, level, is_first=(i == 0), show_answers=show_answers)


def add_success_criteria(doc, criteria_list, theme_key='classic', level='expected'):
    """Add a themed success criteria checklist."""
    theme = THEMES[theme_key]
    is_dev = level == 'developing'

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(10)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLOURS['criteria_bg'])
    set_cell_borders(cell, COLOURS['criteria_border'], sz=10)
    set_cell_padding(cell, top=100, bottom=100, left=150, right=150)

    # Title
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{theme['icon']} {theme['criteria']}")
    set_run_font(run, size=Pt(16) if is_dev else Pt(14), bold=True, colour=COLOURS['criteria_text'])

    # Checklist items
    for item in criteria_list:
        pi = cell.add_paragraph()
        set_no_spacing(pi)
        pi.paragraph_format.space_before = Pt(4) if is_dev else Pt(3)
        pi.paragraph_format.space_after = Pt(4) if is_dev else Pt(3)
        run = pi.add_run(f'\u2610  {item}')
        set_run_font(run, size=Pt(14) if is_dev else Pt(12), colour=COLOURS['black'])


def add_eal_glossary_space(doc):
    """Add an empty bordered box for EAL first-language translations."""
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_borders(cell, 'BDBDBD', sz=8)
    set_cell_padding(cell, top=100, bottom=200, left=150, right=150)

    p = cell.paragraphs[0]
    set_no_spacing(p)
    run = p.add_run('My Word Translations / Notes:')
    set_run_font(run, size=Pt(12), bold=True, colour=COLOURS['grey_text'])

    # Add some blank lines for writing
    for _ in range(4):
        blank = cell.add_paragraph()
        set_no_spacing(blank)
        blank.paragraph_format.space_before = Pt(2)
        r = blank.add_run('_' * 70)
        set_run_font(r, size=Pt(11), colour=COLOURS['hint_text'])


def add_footer(doc, level='expected', worksheet_type='Worksheet'):
    """Add a small footer label showing differentiation level."""
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_no_spacing(p)
    label = DIFF_LEVELS[level]['label']
    run = p.add_run(f'{worksheet_type} \u2014 {label}')
    set_run_font(run, size=Pt(9), italic=True, colour=COLOURS['hint_text'])


# ─── Matching Activity Components ──────────────────────────────────────────────


def add_matching_table(doc, pairs, level='expected'):
    """
    Add a matching activity table with two shuffled columns.

    pairs: list of {"left": str, "right": str} dicts
    """
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    import random
    left_items = [p['left'] for p in pairs]
    right_items = [p['right'] for p in pairs]
    random.shuffle(right_items)

    table = doc.add_table(rows=len(pairs) + 1, cols=3)
    set_table_full_width(table)

    # Header row
    headers = ['Term', '', 'Definition']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, 'E0E0E0')
        set_cell_padding(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        run = p.add_run(header)
        set_run_font(run, size=Pt(font_size), bold=True, colour=COLOURS['black'])

    # Data rows
    for row_idx, (left, right) in enumerate(zip(left_items, right_items), start=1):
        # Left column
        cell_l = table.cell(row_idx, 0)
        set_cell_padding(cell_l, top=60, bottom=60, left=80, right=80)
        p_l = cell_l.paragraphs[0]
        set_no_spacing(p_l)
        run_l = p_l.add_run(left)
        set_run_font(run_l, size=Pt(font_size), colour=COLOURS['black'])

        # Arrow column
        cell_m = table.cell(row_idx, 1)
        set_cell_padding(cell_m, top=60, bottom=60, left=40, right=40)
        p_m = cell_m.paragraphs[0]
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p_m)
        run_m = p_m.add_run('\u2192')
        set_run_font(run_m, size=Pt(font_size + 2), colour=COLOURS['hint_text'])

        # Right column
        cell_r = table.cell(row_idx, 2)
        set_cell_padding(cell_r, top=60, bottom=60, left=80, right=80)
        p_r = cell_r.paragraphs[0]
        set_no_spacing(p_r)
        run_r = p_r.add_run(right)
        set_run_font(run_r, size=Pt(font_size), colour=COLOURS['black'])


# ─── Sentence Builder Components ───────────────────────────────────────────────


def _add_word_card_row(doc, parts, font_size):
    """Add a single row of colour-coded word cards."""
    if not parts:
        return
    table = doc.add_table(rows=1, cols=len(parts))
    set_table_full_width(table)
    remove_table_borders(table)

    tblPr = table._tbl.tblPr
    cell_spacing = parse_xml(f'<w:tblCellSpacing {nsdecls("w")} w:w="30" w:type="dxa"/>')
    tblPr.append(cell_spacing)

    for i, part in enumerate(parts):
        wt = WORD_TYPES.get(part.get('word_type', 'noun'), WORD_TYPES['noun'])
        cell = table.cell(0, i)
        set_cell_shading(cell, wt['bg'])
        set_cell_borders(cell, wt['border'], sz=8)
        set_cell_padding(cell, top=60, bottom=60, left=60, right=60)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        run = p.add_run(part['part'])
        set_run_font(run, size=Pt(font_size), bold=True, colour=wt['text'])


def add_sentence_builder_box(doc, sentence_parts, level='expected'):
    """
    Add a sentence builder activity with word/phrase cards.

    sentence_parts: list of {"part": str, "word_type": str} dicts
    """
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    import random
    shuffled = sentence_parts[:]
    random.shuffle(shuffled)

    # Render word cards in rows of up to 5
    for row_start in range(0, len(shuffled), 5):
        row_parts = shuffled[row_start:row_start + 5]
        _add_word_card_row(doc, row_parts, font_size)

    # Writing line
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(4)

    p_write = doc.add_paragraph()
    set_no_spacing(p_write)
    r = p_write.add_run('Write your sentence: ')
    set_run_font(r, size=Pt(font_size), bold=True, colour=COLOURS['grey_text'])

    p_line = doc.add_paragraph()
    set_no_spacing(p_line)
    p_line.paragraph_format.space_before = Pt(4)
    r_line = p_line.add_run('_' * 70)
    set_run_font(r_line, size=Pt(font_size), colour=COLOURS['hint_text'])


# ─── Answer Key Components ────────────────────────────────────────────────────


def add_matching_answer_table(doc, pairs, level='expected'):
    """Display correct matching pairs in order (NOT shuffled) for the answer key."""
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    table = doc.add_table(rows=len(pairs) + 1, cols=3)
    set_table_full_width(table)

    # Header row (green tinted for answer key)
    headers = ['Term', '', 'Correct Match']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, 'C8E6C9')
        set_cell_padding(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        run = p.add_run(header)
        set_run_font(run, size=Pt(font_size), bold=True, colour=COLOURS['criteria_text'])

    # Data rows — correct pairs, not shuffled
    for row_idx, pair in enumerate(pairs, start=1):
        cell_l = table.cell(row_idx, 0)
        set_cell_padding(cell_l, top=60, bottom=60, left=80, right=80)
        p_l = cell_l.paragraphs[0]
        set_no_spacing(p_l)
        run_l = p_l.add_run(f"{row_idx}. {pair['left']}")
        set_run_font(run_l, size=Pt(font_size), colour=COLOURS['black'])

        cell_m = table.cell(row_idx, 1)
        set_cell_padding(cell_m, top=60, bottom=60, left=40, right=40)
        p_m = cell_m.paragraphs[0]
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p_m)
        run_m = p_m.add_run('\u2192')
        set_run_font(run_m, size=Pt(font_size + 2), colour=COLOURS['criteria_text'])

        cell_r = table.cell(row_idx, 2)
        set_cell_padding(cell_r, top=60, bottom=60, left=80, right=80)
        p_r = cell_r.paragraphs[0]
        set_no_spacing(p_r)
        run_r = p_r.add_run(pair['right'])
        set_run_font(run_r, size=Pt(font_size), bold=True, colour=COLOURS['criteria_text'])


def add_answer_sentence(doc, sentence_text, level='expected'):
    """Display the correct sentence for a sentence builder answer key."""
    diff = DIFF_LEVELS[level]
    p = doc.add_paragraph()
    set_no_spacing(p)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(f'Answer: {sentence_text}')
    set_run_font(run, size=Pt(diff['font_size']), bold=True, colour=COLOURS['criteria_text'])


# ─── Reading Comprehension Components ─────────────────────────────────────────


def add_reading_passage(doc, passage_data, theme_key='classic', level='expected'):
    """Add a reading passage in a themed container box."""
    theme = THEMES[theme_key]
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'

    # Passage title above the container
    passage_title = passage_data.get('title', '')
    if passage_title:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(8)
        spacer.paragraph_format.space_after = Pt(2)

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p_title)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run(f'{theme["icon"]} {passage_title}')
        set_run_font(run_title, size=Pt(18) if is_dev else Pt(16), bold=True, colour=COLOURS['title_text'])

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, theme['body'])
    padding = diff['padding']
    set_cell_padding(cell, top=padding, bottom=padding, left=200, right=200)

    # Split passage into paragraphs on double newlines
    text = passage_data.get('text', '')
    paragraphs = text.split('\n\n') if '\n\n' in text else [text]

    for i, para_text in enumerate(paragraphs):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        set_no_spacing(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(diff['line_spacing'])
        run = p.add_run(para_text.strip())
        set_run_font(run, size=Pt(diff['font_size']), colour=COLOURS['black'])

    # Source note if present
    if passage_data.get('source_note'):
        p_src = cell.add_paragraph()
        set_no_spacing(p_src)
        p_src.paragraph_format.space_before = Pt(8)
        run_src = p_src.add_run(passage_data['source_note'])
        set_run_font(run_src, size=Pt(9), italic=True, colour=COLOURS['hint_text'])


def add_vocabulary_box(doc, vocabulary, level='expected'):
    """Add a vocabulary key box listing key words from the passage."""
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLOURS['reminder_bg'])
    set_cell_borders(cell, COLOURS['reminder_border'], sz=6)
    set_cell_padding(cell, top=80, bottom=80, left=150, right=150)

    # Header
    p = cell.paragraphs[0]
    set_no_spacing(p)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Key Vocabulary')
    set_run_font(run, size=Pt(14) if is_dev else Pt(12), bold=True, colour=COLOURS['reminder_text'])

    # Each vocabulary word
    for item in vocabulary:
        wt = WORD_TYPES.get(item.get('word_type', 'noun'), WORD_TYPES['noun'])
        pw = cell.add_paragraph()
        set_no_spacing(pw)
        pw.paragraph_format.space_before = Pt(2)
        pw.paragraph_format.space_after = Pt(2)

        r1 = pw.add_run(f'{wt["symbol"]} {item.get("word", "")}')
        set_run_font(r1, size=Pt(diff['font_size'] - 2), bold=True, colour=wt['text'])
        r2 = pw.add_run(f' \u2014 {item.get("definition", "")}')
        set_run_font(r2, size=Pt(diff['font_size'] - 3), italic=True, colour=COLOURS['grey_text'])


def add_comprehension_questions(doc, questions, level='expected', show_answers=False):
    """Add numbered comprehension questions with type badges and answer spaces."""
    from docx.shared import RGBColor
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'
    font_size = diff['font_size']

    # Question type visual coding
    TYPE_LABELS = {
        'retrieval': ('Find and Copy', 'E8F5E9', '2E7D32'),
        'inference': ('Think and Infer', 'E3F2FD', '1565C0'),
        'vocabulary': ('Word Meaning', 'FFF8E1', 'F57F17'),
        'author_intent': ("Author's Choice", 'F3E5F5', '7B1FA2'),
        'evaluation': ('Your Opinion', 'FCE4EC', 'C62828'),
    }

    for q in questions:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(2)

        # Question type badge
        q_type = q.get('question_type', 'retrieval')
        label, bg_hex, text_hex = TYPE_LABELS.get(q_type, TYPE_LABELS['retrieval'])

        badge_table = doc.add_table(rows=1, cols=1)
        set_table_full_width(badge_table)
        remove_table_borders(badge_table)
        badge_cell = badge_table.cell(0, 0)
        set_cell_shading(badge_cell, bg_hex)
        set_cell_borders(badge_cell, text_hex, sz=4)
        set_cell_padding(badge_cell, top=40, bottom=40, left=100, right=100)

        p_badge = badge_cell.paragraphs[0]
        set_no_spacing(p_badge)
        r_type = p_badge.add_run(f'{label}  ')
        set_run_font(r_type, size=Pt(9), bold=True, colour=RGBColor.from_string(text_hex))
        marks = q.get('marks', 1)
        r_marks = p_badge.add_run(f'[{marks} mark{"s" if marks > 1 else ""}]')
        set_run_font(r_marks, size=Pt(9), italic=True, colour=COLOURS['hint_text'])

        # Question text
        p_q = doc.add_paragraph()
        set_no_spacing(p_q)
        p_q.paragraph_format.space_before = Pt(4)
        p_q.paragraph_format.space_after = Pt(4)
        run_q = p_q.add_run(f'{q.get("number", "?")}. {q.get("question", "")}')
        set_run_font(run_q, size=Pt(font_size), bold=True, colour=COLOURS['black'])

        # Word bank hints for developing level
        if is_dev and q.get('word_bank'):
            p_wb = doc.add_paragraph()
            set_no_spacing(p_wb)
            p_wb.paragraph_format.space_after = Pt(4)
            hint_text = 'Hint words: ' + ', '.join(q['word_bank'])
            run_wb = p_wb.add_run(hint_text)
            set_run_font(run_wb, size=Pt(font_size - 2), italic=True, colour=COLOURS['hint_text'])

        # Answer area
        if show_answers:
            p_ans = doc.add_paragraph()
            set_no_spacing(p_ans)
            p_ans.paragraph_format.space_before = Pt(4)
            run_ans = p_ans.add_run(f'Answer: {q.get("answer", "[answer not provided]")}')
            set_run_font(run_ans, size=Pt(font_size - 1), bold=True, colour=COLOURS['criteria_text'])
        else:
            num_lines = q.get('lines', 2)
            for _ in range(num_lines):
                p_line = doc.add_paragraph()
                set_no_spacing(p_line)
                p_line.paragraph_format.space_before = Pt(4)
                p_line.paragraph_format.space_after = Pt(4)
                run_line = p_line.add_run('_' * 70)
                set_run_font(run_line, size=Pt(font_size), colour=COLOURS['hint_text'])
