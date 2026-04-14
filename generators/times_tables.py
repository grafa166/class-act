"""
Times Tables drill worksheet generator.

Produces a themed, pure-equation Word document of multiplication and
division facts. Designed for fast recall practice — no wordy questions,
just a clean grid of number facts across two columns.
"""

import io
import math
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_section_header,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_no_spacing,
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    set_table_full_width,
    remove_table_borders,
)
from generators.styles import COLOURS, DIFF_LEVELS, THEMES
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_times_tables_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a times tables drill worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - sections: list of {title, instructions, tables_focus, facts: [{question, answer}]}
            - speed_challenge: {title, instructions, time_limit_seconds} (optional)
            - success_criteria: list of str
        theme_key: Visual theme key
        level: Differentiation level
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space
        show_answers: Whether to display answers (teacher answer key mode)

    Returns:
        BytesIO buffer containing the .docx file
    """
    diff = DIFF_LEVELS[level]
    theme = THEMES[theme_key]

    # 1. Base document
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Title area
    add_title_area(
        doc,
        content['title'],
        'Answer each times table fact as quickly and accurately as you can!',
        theme_key,
        level,
        is_answer_key=show_answers,
    )

    # 3. Learning objective
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Each section with its fact grid
    for section_number, section in enumerate(content.get('sections', []), start=1):
        add_section_header(doc, section_number, section.get('title', ''), theme_key)

        # Brief instruction
        instructions = section.get('instructions', '')
        if instructions:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_inst = doc.add_paragraph()
            set_no_spacing(p_inst)
            p_inst.paragraph_format.space_after = Pt(6)
            run_inst = p_inst.add_run(instructions)
            set_run_font(
                run_inst,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Build a 2-column fact grid
        facts = section.get('facts', [])
        num_rows = math.ceil(len(facts) / 2)

        if num_rows > 0:
            table = doc.add_table(rows=num_rows, cols=2)
            set_table_full_width(table)
            remove_table_borders(table)

            cell_bg = theme['body']
            cell_border = theme['accent']

            for idx, fact in enumerate(facts):
                row_idx = idx // 2
                col_idx = idx % 2
                cell = table.cell(row_idx, col_idx)

                set_cell_shading(cell, cell_bg)
                set_cell_borders(cell, cell_border, sz=6)
                set_cell_padding(cell, top=100, bottom=100, left=150, right=150)

                # Fact number + question
                p_q = cell.paragraphs[0]
                set_no_spacing(p_q)
                p_q.paragraph_format.space_after = Pt(2)

                run_num = p_q.add_run(f'{idx + 1}.  ')
                set_run_font(
                    run_num,
                    size=Pt(diff['font_size']),
                    bold=True,
                    colour=RGBColor.from_string(theme['header']),
                )

                question_text = fact.get('question', '')
                run_q = p_q.add_run(question_text)
                set_run_font(
                    run_q,
                    size=Pt(diff['font_size'] + 2),  # Larger for visibility
                    bold=True,
                    colour=COLOURS['black'],
                )

                # Answer (answer key only)
                if show_answers and fact.get('answer') is not None:
                    p_ans = cell.add_paragraph()
                    set_no_spacing(p_ans)
                    p_ans.paragraph_format.space_before = Pt(2)
                    run_ans = p_ans.add_run(f'= {fact["answer"]}')
                    set_run_font(
                        run_ans,
                        size=Pt(diff['font_size'] + 1),
                        bold=True,
                        colour=COLOURS['criteria_text'],
                    )

            # Clear leftover empty cell if odd number of facts
            if len(facts) % 2 == 1:
                empty_cell = table.cell(num_rows - 1, 1)
                empty_cell.paragraphs[0].clear()

    # 5. Optional speed challenge
    speed = content.get('speed_challenge')
    if speed and level != 'developing':
        challenge_number = len(content.get('sections', [])) + 1
        add_section_header(doc, challenge_number, speed.get('title', 'Speed Challenge'), theme_key)

        if speed.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_chal = doc.add_paragraph()
            set_no_spacing(p_chal)
            p_chal.paragraph_format.space_after = Pt(6)
            instr_text = speed['instructions']
            if speed.get('time_limit_seconds'):
                instr_text = f"[{speed['time_limit_seconds']} seconds] " + instr_text
            run_chal = p_chal.add_run(instr_text)
            set_run_font(
                run_chal,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Render speed challenge facts (if any) or blank lines
        speed_facts = speed.get('facts', [])
        if speed_facts:
            num_rows = math.ceil(len(speed_facts) / 2)
            table = doc.add_table(rows=num_rows, cols=2)
            set_table_full_width(table)
            remove_table_borders(table)

            for idx, fact in enumerate(speed_facts):
                row_idx = idx // 2
                col_idx = idx % 2
                cell = table.cell(row_idx, col_idx)
                set_cell_shading(cell, theme['body'])
                set_cell_borders(cell, theme['accent'], sz=6)
                set_cell_padding(cell, top=80, bottom=80, left=120, right=120)

                p_q = cell.paragraphs[0]
                set_no_spacing(p_q)
                run_q = p_q.add_run(fact.get('question', ''))
                set_run_font(
                    run_q,
                    size=Pt(diff['font_size'] + 1),
                    bold=True,
                    colour=COLOURS['black'],
                )

                if show_answers and fact.get('answer') is not None:
                    p_ans = cell.add_paragraph()
                    set_no_spacing(p_ans)
                    run_ans = p_ans.add_run(f'= {fact["answer"]}')
                    set_run_font(
                        run_ans,
                        size=Pt(diff['font_size']),
                        bold=True,
                        colour=COLOURS['criteria_text'],
                    )

            if len(speed_facts) % 2 == 1:
                empty_cell = table.cell(num_rows - 1, 1)
                empty_cell.paragraphs[0].clear()

    # 6. Success criteria
    add_success_criteria(doc, content.get('success_criteria', []), theme_key, level)

    # 7. EAL glossary
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 8. Footer
    add_footer(doc, level, 'Times Tables Drill')

    # 9. Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
