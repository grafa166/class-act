"""
Sentence builder worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children arrange colour-coded word cards into complete sentences.
"""

import io
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_colour_key,
    add_sentence_builder_box,
    add_answer_sentence,
    add_section_header,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_no_spacing,
)
from generators.styles import COLOURS, DIFF_LEVELS
from docx.shared import Pt


def generate_sentence_builder_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a sentence builder worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - exercises: list of {title, instructions, sentence_parts: [{part, word_type}]}
            - extension: {title, instructions, lines: int} (optional)
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space

    Returns:
        BytesIO buffer containing the .docx file
    """
    diff = DIFF_LEVELS[level]

    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    add_title_area(
        doc,
        content['title'],
        'Build sentences from the word cards!',
        theme_key,
        level,
        is_answer_key=show_answers,
    )

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Add colour key showing word types
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    # Collect all unique word types from exercises for the key
    word_types_used = []
    for exercise in content['exercises']:
        for part in exercise.get('sentence_parts', []):
            wt = part.get('word_type', 'noun')
            if wt not in word_types_used:
                word_types_used.append(wt)

    # Show the colour key with all word types found in the exercises
    if word_types_used:
        add_colour_key(doc, level, word_types_to_show=word_types_used)
    else:
        add_colour_key(doc, level)

    # 5. Add each exercise: section header, instructions, and sentence builder box
    for exercise_number, exercise in enumerate(content['exercises'], start=1):
        # Section header
        add_section_header(doc, exercise_number, exercise['title'], theme_key)

        # Exercise instruction paragraph
        if exercise.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_inst = doc.add_paragraph()
            set_no_spacing(p_inst)
            p_inst.paragraph_format.space_after = Pt(6)
            run_inst = p_inst.add_run(exercise['instructions'])
            set_run_font(
                run_inst,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Answer key shows correct sentence; student version shows shuffled cards
        if show_answers:
            correct = exercise.get('correct_sentence', '')
            add_answer_sentence(doc, correct, level)
        else:
            add_sentence_builder_box(doc, exercise['sentence_parts'], level)

    # 6. Add extension activity (only for expected and greater_depth)
    extension = content.get('extension')
    if extension and level != 'developing':
        exercise_number = len(content['exercises']) + 1
        add_section_header(doc, exercise_number, extension['title'], theme_key)

        # Extension instructions
        if extension.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_ext = doc.add_paragraph()
            set_no_spacing(p_ext)
            p_ext.paragraph_format.space_after = Pt(6)
            run_ext = p_ext.add_run(extension['instructions'])
            set_run_font(
                run_ext,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Blank writing lines for the extension activity
        num_lines = extension.get('lines', 4)
        for _ in range(num_lines):
            p_line = doc.add_paragraph()
            set_no_spacing(p_line)
            p_line.paragraph_format.space_before = Pt(6)
            p_line.paragraph_format.space_after = Pt(6)
            run_line = p_line.add_run('_' * 70)
            set_run_font(
                run_line,
                size=Pt(diff['font_size']),
                colour=COLOURS['hint_text'],
            )

    # 7. Add success criteria checklist
    add_success_criteria(doc, content['success_criteria'], theme_key, level)

    # 8. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 9. Add footer with differentiation level label
    add_footer(doc, level, 'Sentence Builder')

    # 10. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
