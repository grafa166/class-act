"""
Maths problem solving worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children read a real-world scenario, interpret data, and answer tiered
problem-solving questions (calculate, explain, estimate, prove).
"""

import io
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_reading_passage,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_table_full_width,
    remove_table_borders,
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    set_no_spacing,
)
from generators.styles import COLOURS, DIFF_LEVELS, THEMES


# Subtitle text varies by differentiation level
_SUBTITLES = {
    'developing': 'Read the problem carefully, then answer the questions!',
    'expected': 'Read the problem and show your working in full!',
    'greater_depth': 'Read the problem closely and explain your reasoning!',
}

# Question type visual coding for maths problem solving
_MATHS_TYPE_LABELS = {
    'calculate': ('Calculate', 'E3F2FD', '1565C0'),
    'explain': ('Explain', 'E8F5E9', '2E7D32'),
    'estimate': ('Estimate', 'FFF8E1', 'F57F17'),
    'prove': ('Prove It', 'F3E5F5', '7B1FA2'),
}


def _add_data_table(doc, data_items, theme_key='classic', level='expected'):
    """
    Add a simple two-column data table showing scenario data.

    data_items: list of {"label": str, "value": str}
    """
    theme = THEMES[theme_key]
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'
    font_size = diff['font_size']

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    # Section label above the table
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p_label)
    p_label.paragraph_format.space_after = Pt(4)
    run_label = p_label.add_run('Key Information')
    set_run_font(run_label, size=Pt(16) if is_dev else Pt(14), bold=True, colour=COLOURS['title_text'])

    # Create the 2-column table: Label | Value
    num_rows = len(data_items) + 1  # +1 for header row
    table = doc.add_table(rows=num_rows, cols=2)
    set_table_full_width(table)

    # Header row
    for col_idx, header_text in enumerate(['Item', 'Value']):
        cell = table.cell(0, col_idx)
        set_cell_shading(cell, theme['header'])
        set_cell_padding(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        run = p.add_run(header_text)
        set_run_font(run, size=Pt(font_size), bold=True, colour=COLOURS['white'])

    # Data rows
    for row_idx, item in enumerate(data_items, start=1):
        # Label cell
        cell_label = table.cell(row_idx, 0)
        set_cell_shading(cell_label, theme['body'])
        set_cell_padding(cell_label, top=50, bottom=50, left=100, right=100)
        p_l = cell_label.paragraphs[0]
        set_no_spacing(p_l)
        run_l = p_l.add_run(item.get('label', ''))
        set_run_font(run_l, size=Pt(font_size), bold=True, colour=COLOURS['black'])

        # Value cell
        cell_value = table.cell(row_idx, 1)
        set_cell_shading(cell_value, theme['body'])
        set_cell_padding(cell_value, top=50, bottom=50, left=100, right=100)
        p_v = cell_value.paragraphs[0]
        set_no_spacing(p_v)
        run_v = p_v.add_run(item.get('value', ''))
        set_run_font(run_v, size=Pt(font_size), colour=COLOURS['black'])


def _add_problem_solving_questions(doc, questions, level='expected', show_answers=False):
    """
    Add numbered problem-solving questions with maths-specific type badges
    and answer spaces.

    Structurally identical to add_comprehension_questions but uses
    maths question type labels (calculate, explain, estimate, prove).
    """
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'
    font_size = diff['font_size']

    for q in questions:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(2)

        # Question type badge
        q_type = q.get('question_type', 'calculate')
        label, bg_hex, text_hex = _MATHS_TYPE_LABELS.get(
            q_type, _MATHS_TYPE_LABELS['calculate']
        )

        badge_table = doc.add_table(rows=1, cols=1)
        set_table_full_width(badge_table)
        remove_table_borders(badge_table)
        badge_cell = badge_table.cell(0, 0)
        set_cell_shading(badge_cell, bg_hex)
        set_cell_borders(badge_cell, text_hex, sz=4)
        set_cell_padding(badge_cell, top=40, bottom=40, left=100, right=100)

        p_badge = badge_cell.paragraphs[0]
        set_no_spacing(p_badge)
        r_type = p_badge.add_run(f'{label}  ')
        set_run_font(r_type, size=Pt(9), bold=True, colour=RGBColor.from_string(text_hex))
        marks = q.get('marks', 1)
        r_marks = p_badge.add_run(f'[{marks} mark{"s" if marks > 1 else ""}]')
        set_run_font(r_marks, size=Pt(9), italic=True, colour=COLOURS['hint_text'])

        # Question text
        p_q = doc.add_paragraph()
        set_no_spacing(p_q)
        p_q.paragraph_format.space_before = Pt(4)
        p_q.paragraph_format.space_after = Pt(4)
        run_q = p_q.add_run(f'{q.get("number", "?")}. {q.get("question", "")}')
        set_run_font(run_q, size=Pt(font_size), bold=True, colour=COLOURS['black'])

        # Word bank hints for developing level
        if is_dev and q.get('word_bank'):
            p_wb = doc.add_paragraph()
            set_no_spacing(p_wb)
            p_wb.paragraph_format.space_after = Pt(4)
            hint_text = 'Hint words: ' + ', '.join(q['word_bank'])
            run_wb = p_wb.add_run(hint_text)
            set_run_font(run_wb, size=Pt(font_size - 2), italic=True, colour=COLOURS['hint_text'])

        # Answer area
        if show_answers:
            p_ans = doc.add_paragraph()
            set_no_spacing(p_ans)
            p_ans.paragraph_format.space_before = Pt(4)
            run_ans = p_ans.add_run(f'Answer: {q.get("answer", "[answer not provided]")}')
            set_run_font(run_ans, size=Pt(font_size - 1), bold=True, colour=COLOURS['criteria_text'])
        else:
            num_lines = q.get('lines', 2)
            for _ in range(num_lines):
                p_line = doc.add_paragraph()
                set_no_spacing(p_line)
                p_line.paragraph_format.space_before = Pt(4)
                p_line.paragraph_format.space_after = Pt(4)
                run_line = p_line.add_run('_' * 70)
                set_run_font(run_line, size=Pt(font_size), colour=COLOURS['hint_text'])


def generate_problem_solving_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a maths problem solving worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - scenario: {title, text, data: [{label, value}]}
            - questions: list of {number, question, question_type, marks,
                                  lines, answer, word_bank}
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space
        show_answers: Whether to render as answer key (teacher edition)

    Returns:
        BytesIO buffer containing the .docx file
    """
    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    subtitle = _SUBTITLES.get(level, _SUBTITLES['expected'])
    add_title_area(doc, content['title'], subtitle, theme_key, level,
                   is_answer_key=show_answers)

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Add the scenario text (reuse reading passage component)
    scenario = content.get('scenario', {})
    passage_data = {
        'title': scenario.get('title', ''),
        'text': scenario.get('text', ''),
    }
    add_reading_passage(doc, passage_data, theme_key, level)

    # 5. Add the data table showing scenario data items
    data_items = scenario.get('data', [])
    if data_items:
        _add_data_table(doc, data_items, theme_key, level)

    # 6. Add problem-solving questions (with maths type badges)
    _add_problem_solving_questions(
        doc, content['questions'], level, show_answers=show_answers
    )

    # 7. Add success criteria checklist
    add_success_criteria(doc, content['success_criteria'], theme_key, level)

    # 8. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 9. Add footer with differentiation level label
    add_footer(doc, level, 'Problem Solving')

    # 10. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
