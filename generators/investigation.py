"""
Science investigation planner worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children plan and record a scientific investigation with variables,
method steps, results table, and conclusion prompts.
"""

import io
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    set_table_full_width,
    remove_table_borders,
    set_no_spacing,
)
from generators.styles import FONT_NAME, COLOURS, THEMES, DIFF_LEVELS


# Subtitle text varies by differentiation level
_SUBTITLES = {
    'developing': 'Follow the steps to plan your investigation!',
    'expected': 'Plan your investigation carefully and record your results!',
    'greater_depth': 'Design a fair test and explain your findings!',
}


def _add_investigation_section_header(doc, title, theme_key='classic'):
    """Add a themed section header bar (without a section number)."""
    theme = THEMES[theme_key]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(10)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, theme['header'])
    set_cell_borders(cell, theme['header'], sz=4)
    set_cell_padding(cell, top=80, bottom=80, left=200, right=200)

    full_title = f"{theme['icon']} {title}"

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p)
    run = p.add_run(full_title)
    set_run_font(run, size=Pt(18), bold=True, colour=COLOURS['white'])


def _add_question_box(doc, question_text, theme_key='classic'):
    """Add the investigation question in a styled reminder-style box."""
    theme = THEMES[theme_key]

    table = doc.add_table(rows=1, cols=1)
    set_table_full_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, theme['body'])
    set_cell_borders(cell, theme['accent'], sz=8)
    set_cell_padding(cell, top=100, bottom=100, left=200, right=200)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_no_spacing(p)
    run = p.add_run(question_text)
    set_run_font(run, size=Pt(16), bold=True, colour=COLOURS['title_text'])


def _add_prediction_section(doc, prediction, prediction_choices, level, theme_key='classic'):
    """Add the prediction section with choices (developing) or sentence starter."""
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'
    font_size = diff['font_size']

    if prediction_choices:
        # Developing level: show multiple-choice options
        for i, choice in enumerate(prediction_choices):
            letter = chr(ord('a') + i)
            p = doc.add_paragraph()
            set_no_spacing(p)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f'    {letter})  {choice}')
            set_run_font(run, size=Pt(font_size), colour=COLOURS['black'])
    else:
        # Expected / Greater Depth: show sentence starter with writing lines
        p = doc.add_paragraph()
        set_no_spacing(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(prediction)
        set_run_font(run, size=Pt(font_size), italic=True, colour=COLOURS['grey_text'])

        # Blank lines for writing
        num_lines = 3 if is_dev else 2
        for _ in range(num_lines):
            p_line = doc.add_paragraph()
            set_no_spacing(p_line)
            p_line.paragraph_format.space_before = Pt(4)
            p_line.paragraph_format.space_after = Pt(4)
            run_line = p_line.add_run('_' * 70)
            set_run_font(run_line, size=Pt(font_size), colour=COLOURS['hint_text'])


def _add_variables_table(doc, variables, level, theme_key='classic'):
    """Add a 3-column fair test variables table."""
    theme = THEMES[theme_key]
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    table = doc.add_table(rows=2, cols=3)
    set_table_full_width(table)

    # Column headings
    headings = ['What we will change', 'What we will measure', 'What we will keep the same']
    for i, heading in enumerate(headings):
        cell = table.cell(0, i)
        set_cell_shading(cell, theme['header'])
        set_cell_borders(cell, theme['header'], sz=6)
        set_cell_padding(cell, top=60, bottom=60, left=80, right=80)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        run = p.add_run(heading)
        set_run_font(run, size=Pt(font_size - 2), bold=True, colour=COLOURS['white'])

    # Data row
    values = [
        variables.get('change', ''),
        variables.get('measure', ''),
        '\n'.join(variables.get('keep_same', [])),
    ]
    for i, value in enumerate(values):
        cell = table.cell(1, i)
        set_cell_shading(cell, theme['body'])
        set_cell_borders(cell, theme['header'], sz=4)
        set_cell_padding(cell, top=80, bottom=80, left=80, right=80)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)
        run = p.add_run(value)
        set_run_font(run, size=Pt(font_size - 1), colour=COLOURS['black'])


def _add_equipment_list(doc, equipment, level):
    """Add equipment as bullet points."""
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    for item in equipment:
        p = doc.add_paragraph()
        set_no_spacing(p)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'\u2022  {item}')
        set_run_font(run, size=Pt(font_size), colour=COLOURS['black'])


def _add_method_steps(doc, method, level):
    """Add numbered method steps."""
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    for i, step in enumerate(method, start=1):
        p = doc.add_paragraph()
        set_no_spacing(p)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r_num = p.add_run(f'{i}.  ')
        set_run_font(r_num, size=Pt(font_size), bold=True, colour=COLOURS['grey_text'])
        r_text = p.add_run(step)
        set_run_font(r_text, size=Pt(font_size), colour=COLOURS['black'])


def _add_results_table(doc, results_table, level, theme_key='classic'):
    """Add a results table with column headings (with units) and empty data rows."""
    theme = THEMES[theme_key]
    diff = DIFF_LEVELS[level]
    font_size = diff['font_size']

    columns = results_table.get('columns', [])
    units = results_table.get('units', [])
    num_rows = results_table.get('rows', 4)

    # Total rows = 1 header row + num_rows data rows
    table = doc.add_table(rows=num_rows + 1, cols=len(columns))
    set_table_full_width(table)

    # Header row with column names and units
    for col_idx, col_name in enumerate(columns):
        cell = table.cell(0, col_idx)
        set_cell_shading(cell, theme['header'])
        set_cell_borders(cell, theme['header'], sz=6)
        set_cell_padding(cell, top=60, bottom=60, left=80, right=80)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_no_spacing(p)

        # Column name
        run = p.add_run(col_name)
        set_run_font(run, size=Pt(font_size - 2), bold=True, colour=COLOURS['white'])

        # Add unit on a new line if provided
        unit = units[col_idx] if col_idx < len(units) else ''
        if unit:
            r_unit = p.add_run(f'\n({unit})')
            set_run_font(r_unit, size=Pt(font_size - 4), italic=True, colour=COLOURS['white'])

    # Empty data rows for pupils to fill in
    for row_idx in range(1, num_rows + 1):
        for col_idx in range(len(columns)):
            cell = table.cell(row_idx, col_idx)
            set_cell_borders(cell, theme['header'], sz=4)
            set_cell_padding(cell, top=60, bottom=60, left=80, right=80)

            # Alternate row shading for readability
            if row_idx % 2 == 0:
                set_cell_shading(cell, theme['body'])

            # Clear cell content (leave empty for writing)
            p = cell.paragraphs[0]
            set_no_spacing(p)
            # Add a non-breaking space to preserve row height
            run = p.add_run('\u00A0')
            set_run_font(run, size=Pt(font_size), colour=COLOURS['black'])


def _add_conclusion_prompts(doc, conclusion_prompts, level):
    """Add conclusion sentence starters with blank writing lines."""
    diff = DIFF_LEVELS[level]
    is_dev = level == 'developing'
    font_size = diff['font_size']
    num_lines = 3 if is_dev else 2

    for prompt_text in conclusion_prompts:
        # Sentence starter
        p = doc.add_paragraph()
        set_no_spacing(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(prompt_text)
        set_run_font(run, size=Pt(font_size), bold=True, colour=COLOURS['grey_text'])

        # Blank lines for writing
        for _ in range(num_lines):
            p_line = doc.add_paragraph()
            set_no_spacing(p_line)
            p_line.paragraph_format.space_before = Pt(4)
            p_line.paragraph_format.space_after = Pt(4)
            run_line = p_line.add_run('_' * 70)
            set_run_font(run_line, size=Pt(font_size), colour=COLOURS['hint_text'])


def generate_investigation_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a science investigation planner worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - investigation: {question, prediction, prediction_choices, variables}
            - equipment: list of str
            - method: list of str
            - results_table: {columns, rows, units}
            - conclusion_prompts: list of str
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space
        show_answers: Whether to render as answer key (teacher edition)

    Returns:
        BytesIO buffer containing the .docx file
    """
    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    subtitle = _SUBTITLES.get(level, _SUBTITLES['expected'])
    add_title_area(doc, content['title'], subtitle, theme_key, level,
                   is_answer_key=show_answers)

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Investigation Question section
    investigation = content.get('investigation', {})
    _add_investigation_section_header(doc, 'Our Investigation Question', theme_key)
    _add_question_box(doc, investigation.get('question', ''), theme_key)

    # 5. Prediction section
    _add_investigation_section_header(doc, 'My Prediction', theme_key)
    prediction_choices = investigation.get('prediction_choices')
    _add_prediction_section(
        doc,
        prediction=investigation.get('prediction', ''),
        prediction_choices=prediction_choices,
        level=level,
        theme_key=theme_key,
    )

    # 6. Variables section (Fair Test)
    variables = investigation.get('variables', {})
    _add_investigation_section_header(doc, 'Fair Test Variables', theme_key)
    _add_variables_table(doc, variables, level, theme_key)

    # 7. Equipment section
    equipment = content.get('equipment', [])
    if equipment:
        _add_investigation_section_header(doc, 'Equipment', theme_key)
        _add_equipment_list(doc, equipment, level)

    # 8. Method section
    method = content.get('method', [])
    if method:
        _add_investigation_section_header(doc, 'Method', theme_key)
        _add_method_steps(doc, method, level)

    # 9. Results Table section
    results_table = content.get('results_table', {})
    if results_table.get('columns'):
        _add_investigation_section_header(doc, 'Results', theme_key)
        _add_results_table(doc, results_table, level, theme_key)

    # 10. Conclusion section
    conclusion_prompts = content.get('conclusion_prompts', [])
    if conclusion_prompts:
        _add_investigation_section_header(doc, 'Conclusion', theme_key)
        _add_conclusion_prompts(doc, conclusion_prompts, level)

    # 11. Add success criteria checklist
    add_success_criteria(doc, content.get('success_criteria', []), theme_key, level)

    # 12. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 13. Add footer with differentiation level label
    add_footer(doc, level, 'Investigation Planner')

    # 14. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
