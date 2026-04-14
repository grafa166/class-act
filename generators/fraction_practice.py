"""
Fraction practice worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children practise fractions with proper visual fraction rendering
(Unicode fraction characters and superscript/subscript notation).
"""

import io
import math
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_section_header,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_no_spacing,
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    set_table_full_width,
    remove_table_borders,
)
from generators.styles import COLOURS, DIFF_LEVELS, FONT_NAME, THEMES
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── Unicode Fraction Mapping ────────────────────────────────────────────────
# Maps common fractions to their Unicode single-character equivalents.
# For fractions not in this map, we render numerator⁄denominator style.

UNICODE_FRACTIONS = {
    (1, 2): '\u00BD',     # ½
    (1, 3): '\u2153',     # ⅓
    (2, 3): '\u2154',     # ⅔
    (1, 4): '\u00BC',     # ¼
    (3, 4): '\u00BE',     # ¾
    (1, 5): '\u2155',     # ⅕
    (2, 5): '\u2156',     # ⅖
    (3, 5): '\u2157',     # ⅗
    (4, 5): '\u2158',     # ⅘
    (1, 6): '\u2159',     # ⅙
    (5, 6): '\u215A',     # ⅚
    (1, 7): '\u2150',     # ⅐
    (1, 8): '\u215B',     # ⅛
    (3, 8): '\u215C',     # ⅜
    (5, 8): '\u215D',     # ⅝
    (7, 8): '\u215E',     # ⅞
    (1, 9): '\u2151',     # ⅑
    (1, 10): '\u2152',    # ⅒
}

# Superscript and subscript digit maps for custom fraction rendering
_SUPERSCRIPT = str.maketrans('0123456789', '⁰¹²³⁴⁵⁶⁷⁸⁹')
_SUBSCRIPT = str.maketrans('0123456789', '₀₁₂₃₄₅₆₇₈₉')
_FRACTION_SLASH = '\u2044'  # ⁄ (fraction slash)


def render_fraction_text(numerator, denominator):
    """
    Return a string that visually represents a fraction.

    Prefers Unicode fraction characters when available (½, ¼, ⅓ etc.),
    falls back to superscript-numerator⁄subscript-denominator style.

    Args:
        numerator: int or str — the top number
        denominator: int or str — the bottom number

    Returns:
        A string like '½' or '³⁄₇'
    """
    try:
        num = int(numerator)
        den = int(denominator)
    except (ValueError, TypeError):
        return f'{numerator}/{denominator}'

    # Check for Unicode single-character fraction
    unicode_char = UNICODE_FRACTIONS.get((num, den))
    if unicode_char:
        return unicode_char

    # Fallback: superscript numerator + fraction slash + subscript denominator
    sup = str(num).translate(_SUPERSCRIPT)
    sub = str(den).translate(_SUBSCRIPT)
    return f'{sup}{_FRACTION_SLASH}{sub}'


def _add_fraction_diagram(cell, shaded, total, theme_key='classic'):
    """
    Render a simple shaded-shape fraction diagram inside a cell.

    Creates a nested table of `total` equal-width cells, with the first
    `shaded` cells filled in the theme accent colour. Good for Year 1-3
    visual recognition of fractions like 3/4, 2/5, etc.

    Args:
        cell: The parent docx cell to add the diagram into
        shaded: Number of cells to fill
        total: Total number of cells (the denominator)
        theme_key: Visual theme key for shading colour
    """
    try:
        shaded_int = int(shaded)
        total_int = int(total)
    except (ValueError, TypeError):
        return
    if total_int <= 0 or total_int > 12 or shaded_int < 0 or shaded_int > total_int:
        return

    theme = THEMES.get(theme_key, THEMES['classic'])
    shaded_colour = theme['accent']
    unshaded_colour = 'FFFFFF'
    border_colour = theme['header']

    # Spacer paragraph before the diagram
    p_spacer = cell.add_paragraph()
    set_no_spacing(p_spacer)
    p_spacer.paragraph_format.space_before = Pt(4)
    p_spacer.paragraph_format.space_after = Pt(2)

    # Nested 1-row table with `total` cells
    diagram = cell.add_table(rows=1, cols=total_int)

    for idx in range(total_int):
        diagram_cell = diagram.cell(0, idx)
        fill = shaded_colour if idx < shaded_int else unshaded_colour
        set_cell_shading(diagram_cell, fill)
        set_cell_borders(diagram_cell, border_colour, sz=6)
        set_cell_padding(diagram_cell, top=140, bottom=140, left=60, right=60)
        # Empty content — just the shaded block
        diagram_cell.paragraphs[0].clear()
        p = diagram_cell.paragraphs[0]
        set_no_spacing(p)
        run = p.add_run(' ')
        set_run_font(run, size=Pt(10), colour=COLOURS['black'])


import re

# Match patterns like 3/4, 12/100, etc. but not dates (2024/03/15),
# file paths, or chained slashes. Excludes digit-or-slash on either side.
_FRACTION_PATTERN = re.compile(r'(?<![\d/])(\d{1,3})/(\d{1,3})(?![\d/])')


def _parse_fraction_from_text(text):
    """
    Attempt to detect and replace fraction notation in a text string.

    Looks for patterns like '3/4', '1/2', '7/10' and replaces them
    with proper Unicode fraction rendering.

    Returns the processed text string.
    """
    def replace_match(m):
        num = int(m.group(1))
        den = int(m.group(2))
        if den == 0:
            return m.group(0)  # Don't replace division by zero
        return render_fraction_text(num, den)

    return _FRACTION_PATTERN.sub(replace_match, text)


def _render_exercise_cell(cell, idx, exercise, theme, theme_key, diff, show_answers):
    """Render a single fraction exercise inside a table cell."""
    set_cell_shading(cell, theme['body'])
    set_cell_borders(cell, theme['accent'], sz=6)
    set_cell_padding(cell, top=120, bottom=120, left=150, right=150)

    # Question line: number + question text (fractions rendered)
    p_q = cell.paragraphs[0]
    set_no_spacing(p_q)
    p_q.paragraph_format.space_after = Pt(4)

    run_num = p_q.add_run(f'{idx + 1}. ')
    set_run_font(
        run_num,
        size=Pt(diff['font_size']),
        bold=True,
        colour=RGBColor.from_string(theme['header']),
    )

    question_text = _parse_fraction_from_text(exercise.get('question', ''))
    run_q = p_q.add_run(question_text)
    set_run_font(
        run_q,
        size=Pt(diff['font_size'] + 2),
        bold=True,
        colour=COLOURS['black'],
    )

    # Visual diagram (shade X of Y cells) — for developing level recognition
    diagram = exercise.get('diagram')
    if diagram and not show_answers:
        _add_fraction_diagram(
            cell,
            diagram.get('shaded', 0),
            diagram.get('total', 0),
            theme_key,
        )

    # Visual hint text (sentence describing what to shade, etc.)
    visual_hint = exercise.get('visual_hint')
    if visual_hint and not show_answers:
        p_hint = cell.add_paragraph()
        set_no_spacing(p_hint)
        p_hint.paragraph_format.space_before = Pt(4)
        run_hint = p_hint.add_run(_parse_fraction_from_text(visual_hint))
        set_run_font(
            run_hint,
            size=Pt(diff['font_size'] - 3),
            italic=True,
            colour=COLOURS['hint_text'],
        )

    # Answer or blank line
    if show_answers and exercise.get('answer'):
        answer_text = _parse_fraction_from_text(str(exercise['answer']))
        p_ans = cell.add_paragraph()
        set_no_spacing(p_ans)
        p_ans.paragraph_format.space_before = Pt(4)
        run_ans = p_ans.add_run(f'Answer: {answer_text}')
        set_run_font(
            run_ans,
            size=Pt(diff['font_size']),
            bold=True,
            colour=COLOURS['criteria_text'],
        )
    elif not show_answers:
        p_ans = cell.add_paragraph()
        set_no_spacing(p_ans)
        p_ans.paragraph_format.space_before = Pt(8)
        run_ans = p_ans.add_run('Answer: _______________')
        set_run_font(
            run_ans,
            size=Pt(diff['font_size']),
            colour=COLOURS['hint_text'],
        )


def _render_exercise_grid(doc, exercises, theme, theme_key, diff, show_answers):
    """Render a 2-column grid of fraction exercises."""
    num_rows = math.ceil(len(exercises) / 2)
    if num_rows == 0:
        return

    table = doc.add_table(rows=num_rows, cols=2)
    set_table_full_width(table)
    remove_table_borders(table)

    for idx, exercise in enumerate(exercises):
        cell = table.cell(idx // 2, idx % 2)
        _render_exercise_cell(
            cell, idx, exercise, theme, theme_key, diff, show_answers,
        )

    # Clear any leftover empty cell when odd number of exercises
    if len(exercises) % 2 == 1:
        empty_cell = table.cell(num_rows - 1, 1)
        empty_cell.paragraphs[0].clear()


def generate_fraction_practice_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a fraction practice worksheet as a Word document.

    All fractions are rendered using proper Unicode fraction characters
    (½, ¼, ⅓) or superscript/subscript notation (⁷⁄₁₂).

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - sections: list of {title, instructions, type, exercises: [{question, answer, visual_hint}]}
            - challenge: {title, instructions, lines: int} (optional)
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space
        show_answers: Whether to display answers (teacher answer key mode)

    Returns:
        BytesIO buffer containing the .docx file
    """
    diff = DIFF_LEVELS[level]
    theme = THEMES[theme_key]

    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    add_title_area(
        doc,
        _parse_fraction_from_text(content['title']),
        'Work with fractions carefully — show your working!',
        theme_key,
        level,
        is_answer_key=show_answers,
    )

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, _parse_fraction_from_text(objective), theme_key)

    # 4. Add each section
    for section_number, section in enumerate(content.get('sections', []), start=1):
        section_title = _parse_fraction_from_text(section.get('title', ''))
        add_section_header(doc, section_number, section_title, theme_key)

        # Section instruction paragraph
        instructions = section.get('instructions', '')
        if instructions:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_inst = doc.add_paragraph()
            set_no_spacing(p_inst)
            p_inst.paragraph_format.space_after = Pt(6)
            run_inst = p_inst.add_run(_parse_fraction_from_text(instructions))
            set_run_font(
                run_inst,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Build the exercise grid (2-column layout)
        _render_exercise_grid(
            doc, section.get('exercises', []), theme, theme_key, diff, show_answers,
        )

    # 5. Add challenge section if present (skip for developing level)
    challenge = content.get('challenge')
    if challenge and level != 'developing':
        challenge_number = len(content.get('sections', [])) + 1
        add_section_header(doc, challenge_number, challenge['title'], theme_key)

        if challenge.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_chal = doc.add_paragraph()
            set_no_spacing(p_chal)
            p_chal.paragraph_format.space_after = Pt(6)
            run_chal = p_chal.add_run(
                _parse_fraction_from_text(challenge['instructions'])
            )
            set_run_font(
                run_chal,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        num_lines = challenge.get('lines', 3)
        for _ in range(num_lines):
            p_line = doc.add_paragraph()
            set_no_spacing(p_line)
            p_line.paragraph_format.space_before = Pt(6)
            p_line.paragraph_format.space_after = Pt(6)
            run_line = p_line.add_run('_' * 70)
            set_run_font(
                run_line,
                size=Pt(diff['font_size']),
                colour=COLOURS['hint_text'],
            )

    # 6. Add success criteria checklist
    add_success_criteria(doc, content.get('success_criteria', []), theme_key, level)

    # 7. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 8. Add footer with differentiation level label
    add_footer(doc, level, 'Fraction Practice')

    # 9. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
