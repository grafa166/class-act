"""
Matching activity worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children draw lines to match related pairs (e.g. terms to definitions).
"""

import io
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_matching_table,
    add_matching_answer_table,
    add_section_header,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_no_spacing,
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    set_table_full_width,
    remove_table_borders,
)
from generators.styles import COLOURS, DIFF_LEVELS, FONT_NAME
from docx.shared import Pt


def generate_matching_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a matching activity worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - activities: list of {title, instructions, pairs: [{left, right}]}
            - bonus_activity: {title, instructions, lines: int} (optional)
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space

    Returns:
        BytesIO buffer containing the .docx file
    """
    diff = DIFF_LEVELS[level]

    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    add_title_area(
        doc,
        content['title'],
        'Draw lines to match the pairs!',
        theme_key,
        level,
        is_answer_key=show_answers,
    )

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Add each matching activity: section header, instructions, and matching table
    for activity_number, activity in enumerate(content['activities'], start=1):
        # Section header
        add_section_header(doc, activity_number, activity['title'], theme_key)

        # Activity instruction paragraph
        if activity.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_inst = doc.add_paragraph()
            set_no_spacing(p_inst)
            p_inst.paragraph_format.space_after = Pt(6)
            run_inst = p_inst.add_run(activity['instructions'])
            set_run_font(
                run_inst,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Matching table — answer key shows correct pairs, student version shuffles
        if show_answers:
            add_matching_answer_table(doc, activity['pairs'], level)
        else:
            add_matching_table(doc, activity['pairs'], level)

    # 5. Add bonus activity section (only for expected and greater_depth)
    bonus = content.get('bonus_activity')
    if bonus and level != 'developing':
        activity_number = len(content['activities']) + 1
        add_section_header(doc, activity_number, bonus['title'], theme_key)

        # Bonus instructions
        if bonus.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_bonus = doc.add_paragraph()
            set_no_spacing(p_bonus)
            p_bonus.paragraph_format.space_after = Pt(6)
            run_bonus = p_bonus.add_run(bonus['instructions'])
            set_run_font(
                run_bonus,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Blank writing lines for the bonus activity
        num_lines = bonus.get('lines', 4)
        for _ in range(num_lines):
            p_line = doc.add_paragraph()
            set_no_spacing(p_line)
            p_line.paragraph_format.space_before = Pt(6)
            p_line.paragraph_format.space_after = Pt(6)
            run_line = p_line.add_run('_' * 70)
            set_run_font(
                run_line,
                size=Pt(diff['font_size']),
                colour=COLOURS['hint_text'],
            )

    # 6. Add success criteria checklist
    add_success_criteria(doc, content['success_criteria'], theme_key, level)

    # 7. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 8. Add footer with differentiation level label
    add_footer(doc, level, 'Matching Activity')

    # 9. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
