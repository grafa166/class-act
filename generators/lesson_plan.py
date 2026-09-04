"""The lesson plan as a Word document she can teach from and edit.

The standard is the one the first prototype failed: *"the bullet point outline
of what the lesson should include — it doesn't speak to what would actually
happen in the lesson."* So every step carries what is on the board at that
moment, the words to say, the questions with the answers to expect, what the
children do, the common wrong answer and how to respond to it, and where the
other adult is. Deep enough to build slides from.

Three decisions govern how it is built, all taken 2026-09-01.

**Boxes are paragraphs, never one-cell tables.** Consecutive paragraphs
carrying identical borders are merged by Word into one visual box, so a box can
hold several paragraphs and still be ordinary, typeable, reflowing text.
`prototype/make_editable_demo.py` builds both versions side by side: in the
table version, pressing Enter inside a box fights the table, a long sentence
stretches a cell rather than reflowing, and select-all-and-change-the-font does
nothing. She opens this the night before she teaches it.

**The formatting lives in named styles.** Every paragraph gets one, and nothing
is formatted on the run. That is what makes changing the font one click, what
lets the Styles pane restyle every objective box at once, and what lets a
paragraph pasted into another document adopt that document's look instead of
dragging Class Act's formatting across with it.

**Arial, black and blue.** The six-colour word-type scheme is retired; the
symbols and labels that carried the meaning stay, so dual-coding survives — and
now survives a mono photocopier and a colour-blind child too, which the colours
never did. Comic Sans is the worksheet generators' default and must not leak in
through a shared helper, which is why nothing here imports `FONT_NAME`. The
joined font used elsewhere in school is out by name.

Nothing here checks whether the teaching is any good. The one thing it refuses
is a worksheet built from a different lesson, because a plan that prints
another lesson's sheet as this lesson's evidence is wrong in a way she would
have to teach to discover. Everything else it renders and labels *"AI-drafted —
check before teaching"*.
"""

import io

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Deliberately not `generators.styles.FONT_NAME`, which is Comic Sans and is
# read by nine call sites in the worksheet generators. Importing it here would
# make this document change the day that one does.
LESSON_PLAN_FONT = "Arial"

# Every style this file defines starts with it, so a paragraph carrying direct
# formatting instead of a style is visible at a glance and in a test.
STYLE_PREFIX = "CA "

# Black and blue, and nothing else. Kept as one set so the test that holds the
# decision has something to check against rather than a list of its own.
BLUE = RGBColor(0x14, 0x44, 0x8C)
BLUE_HEX = "14448C"
BLUE_FILL = "F2F6FC"
RULE_HEX = "C9D2DC"
BLACK = RGBColor(0x00, 0x00, 0x00)

ALLOWED_COLOURS = frozenset({"14448C", "F2F6FC", "C9D2DC", "000000", "FFFFFF"})

# The three vocabulary bands, in the order of difficulty they are, with the
# label saying who each is for. Her words: *"hard, soft and rough are too easy
# for some and too difficult for others."*
VOCABULARY_BANDS = (
    ("everyone", "Everyone leaves with", "●"),
    ("expected", "Expected of most", "◆"),
    ("stretch", "Stretch — offered to all", "★"),
)

# Labels carry what the retired colour scheme used to. A symbol survives a mono
# photocopier; a green box does not.
ADAPTATION_LABELS = (
    ("eal", "EAL — same objective, different route"),
    ("send", "SEND — same objective, different route"),
    ("stretch", "Stretch"),
)


class LessonPlanError(ValueError):
    """The document would say something that is not true of this lesson."""


def _border(edge, colour, sz=8, space=6):
    element = OxmlElement(f"w:{edge}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(sz))
    element.set(qn("w:space"), str(space))
    element.set(qn("w:color"), colour)
    return element


def _box(paragraph, colour=BLUE_HEX, fill=None, sz=8, edges=("top", "left", "bottom", "right")):
    """Draw a box around a PARAGRAPH.

    The whole editability decision in one function. Word merges consecutive
    paragraphs carrying identical borders into a single visual box, so a box
    can hold a heading and three sentences and still be text she can type in.
    """
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in edges:
        borders.append(_border(edge, colour, sz))
    properties.append(borders)
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)


def _build_styles(doc, font):
    """Define every style the document uses, once.

    `Normal` is set as well as the named styles, because a paragraph she adds
    herself inherits from it — a document whose own styles are Arial and whose
    Normal is something else is a document that changes font the moment she
    presses Enter at the end.
    """
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    normal.paragraph_format.space_after = Pt(4)

    def new(name, size, bold=False, colour=None, after=4, before=0, indent=None,
            italic=False, caps=False):
        style = doc.styles.add_style(STYLE_PREFIX + name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.all_caps = caps
        style.font.color.rgb = colour if colour is not None else BLACK
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.space_before = Pt(before)
        if indent is not None:
            style.paragraph_format.left_indent = Cm(indent)
        return style

    new("Title", 16, bold=True, colour=BLUE, after=2)
    new("Subtitle", 10, colour=BLUE, after=10)
    new("Section", 12, bold=True, colour=BLUE, after=4, before=12)
    # One style, one meaning. Restyling every objective box at once from the
    # Styles pane is half the reason the styles exist, and a style named
    # Objective that also carried the misconceptions and the worksheet's
    # criteria would restyle those too, silently. `Item` heads a small block;
    # `Tick` is a checklist she works down.
    new("Objective", 12, bold=True, after=4)
    new("Item", 12, bold=True, after=2, before=8)
    new("Criterion", 11, indent=0.7, after=2)
    new("Tick", 11, indent=0.7, after=2)
    new("Body", 11, after=4)
    new("Detail", 11, indent=0.7, after=2)
    # `all_caps` rather than upper-casing the text, so what the file holds is
    # still the words she wrote and a test can find them.
    new("Tag", 8, bold=True, colour=BLUE, after=1, before=6, caps=True)
    new("StepHeading", 12, bold=True, colour=BLUE, after=2, before=10)
    new("Note", 10, italic=True, colour=BLUE, after=4, before=8)


def _say(doc, text, style, box=None, fill=None, align=None):
    paragraph = doc.add_paragraph(str(text), style=STYLE_PREFIX + style)
    if box:
        _box(paragraph, colour=box, fill=fill)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def _labelled(doc, label, text, style="Detail"):
    """A label the reader can scan for, and the text after it, in one paragraph.

    One paragraph rather than two, because two would let Word break a page
    between the label and the thing it labels.
    """
    paragraph = doc.add_paragraph(style=STYLE_PREFIX + style)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(text))
    return paragraph


def build_lesson_plan(
    lesson,
    unit_title="",
    subject="",
    year_group="",
    lesson_minutes=60,
    lesson_count=None,
    anchor="",
    outcome="",
    worksheet=None,
    font=LESSON_PLAN_FONT,
):
    """The whole plan, in the order she reads it.

    Raises:
        LessonPlanError: the worksheet was built from a different lesson.
    """
    if worksheet is not None and worksheet.objective.strip() != lesson.objective.strip():
        # Not a formatting concern. A plan that prints another lesson's sheet
        # as this lesson's evidence is wrong in a way that only shows up in
        # the room.
        raise LessonPlanError(
            "That worksheet was built for a different objective, so it cannot "
            "be printed as this lesson's evidence.\n"
            f"  The lesson says: {lesson.objective}\n"
            f"  The sheet says:  {worksheet.objective}"
        )

    doc = Document()
    _build_styles(doc, font)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.6)
        section.left_margin = section.right_margin = Cm(1.8)
    if doc.paragraphs:
        first = doc.paragraphs[0]
        first._element.getparent().remove(first._element)

    _header(doc, lesson, unit_title, subject, year_group, lesson_minutes,
            lesson_count, anchor, outcome)
    _objective_and_criteria(doc, lesson)
    _vocabulary(doc, lesson)
    _steps(doc, lesson)
    _access(doc, lesson)
    _misconceptions(doc, lesson)
    _assessment(doc, lesson)
    _resources(doc, lesson)
    if worksheet is not None:
        _the_worksheet(doc, worksheet)
    _next_lesson(doc, lesson)
    _say(doc, lesson.source, "Note")
    return doc


def _header(doc, lesson, unit_title, subject, year_group, lesson_minutes,
            lesson_count, anchor, outcome):
    where = f"Lesson {lesson.number}" if lesson.number else "Lesson"
    if lesson.number and lesson_count:
        where = f"Lesson {lesson.number} of {lesson_count}"
    _say(doc, unit_title or "Lesson plan", "Title")
    _say(
        doc,
        "  ·  ".join(
            part for part in (where, subject, year_group, f"{lesson_minutes} minutes")
            if part
        ),
        "Subtitle",
    )
    if outcome:
        _labelled(doc, "Unit outcome", outcome, style="Body")
    if anchor:
        # Her evidence to a subject leader that the scheme was followed.
        _labelled(doc, "Where this comes from", anchor, style="Body")
    if lesson.builds_on and lesson.builds_on_reason:
        _labelled(
            doc,
            f"Builds on lesson {lesson.builds_on}",
            lesson.builds_on_reason,
            style="Body",
        )


def _objective_and_criteria(doc, lesson):
    """Together in one box, because they are one thing to her and to a child."""
    _say(doc, "We are learning to", "Tag", box=BLUE_HEX, fill=BLUE_FILL)
    _say(doc, lesson.objective, "Objective", box=BLUE_HEX, fill=BLUE_FILL)
    _say(doc, "I will know I can do it when", "Tag", box=BLUE_HEX, fill=BLUE_FILL)
    for criterion in lesson.success_criteria:
        paragraph = doc.add_paragraph(style=STYLE_PREFIX + "Criterion")
        paragraph.add_run(f"☐  {criterion.criterion}")
        _box(paragraph, colour=BLUE_HEX, fill=BLUE_FILL)
        evidence = doc.add_paragraph(style=STYLE_PREFIX + "Criterion")
        run = evidence.add_run(f"      Look for: {criterion.evidence}")
        run.italic = True
        _box(evidence, colour=BLUE_HEX, fill=BLUE_FILL)


def _vocabulary(doc, lesson):
    _say(doc, "Vocabulary", "Section")
    for band, label, symbol in VOCABULARY_BANDS:
        words = getattr(lesson.vocabulary, band)
        _labelled(doc, f"{symbol} {label}", "  ".join(words))
    # Labelled, because unlabelled it sits under the stretch band and reads
    # as a fourth list of stretch words. Found reading a real plan.
    _labelled(doc, "How to use them", lesson.vocabulary.guidance)


def _steps(doc, lesson):
    _say(doc, "The lesson", "Section")
    for position, step in enumerate(lesson.steps, 1):
        _say(doc, f"{position}. {step.name} — {step.minutes} min", "StepHeading")
        _labelled(doc, "On the board", step.on_the_board)
        _labelled(doc, "Say", step.teacher_says)
        for question in step.questions:
            _labelled(doc, "Ask", question.ask)
            _labelled(doc, "Expect", question.expect)
        _labelled(doc, "Children", step.children_do)
        for watch in step.watch_for:
            _labelled(doc, "Watch for", watch.wrong)
            _labelled(doc, "Respond", watch.respond)
        if step.adults:
            _labelled(doc, "Other adult", step.adults)
        if step.builds_on_step:
            _labelled(doc, "Why this step now", step.builds_on_step)


def _access(doc, lesson):
    """Same objective, different route. Never a different objective."""
    _say(doc, "Reaching the same objective", "Section")
    for key, label in ADAPTATION_LABELS:
        if lesson.adaptations.get(key):
            _labelled(doc, label, lesson.adaptations[key])


def _misconceptions(doc, lesson):
    _say(doc, "What to expect them to get wrong", "Section")
    for misconception in lesson.misconceptions:
        _say(doc, misconception["misconception"], "Item")
        if misconception.get("why"):
            _labelled(doc, "Why", misconception["why"])
        if misconception.get("address"):
            _labelled(doc, "Do", misconception["address"])


def _assessment(doc, lesson):
    _say(doc, "Assessing it", "Section")
    _labelled(doc, "Look for", lesson.assessment.look_for)
    # Asked for by name. "Look for children meeting the criterion" tells her
    # nothing she did not already know.
    _labelled(doc, "Not yet, and why", lesson.assessment.not_yet_example)


def _resources(doc, lesson):
    _say(doc, "Resources", "Section")
    for resource in lesson.resources:
        _say(doc, f"☐  {resource['item']} — {resource['quantity']}", "Tick")


def _the_worksheet(doc, worksheet):
    """Which task on the sheet proves which criterion.

    The headline feature, said on the page she prints. A plan and a sheet that
    each name the other is what makes the pair defensible to a moderator; a
    plan that just says "worksheet attached" is not.
    """
    _say(doc, "The worksheet, and what it proves", "Section")
    title = str(worksheet.content.get("title", "")).strip()
    if title:
        _labelled(doc, "Sheet", title, style="Body")
    for claim in worksheet.evidence:
        _say(doc, claim.criterion, "Item")
        _labelled(doc, "On the sheet", claim.where)
        _labelled(doc, "The task", claim.quote)
        _labelled(doc, "They write", claim.pupil_writes)


def _next_lesson(doc, lesson):
    if lesson.next_lesson:
        _say(doc, "Next lesson", "Section")
        _say(doc, lesson.next_lesson, "Body")


def lesson_plan_bytes(lesson, **kwargs):
    """The document as bytes, ready for a download button."""
    buffer = io.BytesIO()
    build_lesson_plan(lesson, **kwargs).save(buffer)
    return buffer.getvalue()


def lesson_plan_filename(lesson, unit_title=""):
    """What lands in her Downloads folder.

    A folder of files called `download.docx` is a folder of nothing. The unit
    title is hers to type, so it is scrubbed rather than trusted — *"Rocks and
    Soils: term 1"* is a plausible thing to type and carries two characters
    that are a path separator or illegal in a filename on the machines she
    uses.
    """
    number = f"lesson {lesson.number}" if lesson.number else "lesson"
    title = "".join(
        character for character in str(unit_title).strip()
        if character.isalnum() or character in " -_&'"
    ).strip()
    title = " ".join(title.split())
    if title:
        return f"{title} - {number} - plan.docx"
    return f"{number.capitalize()} - plan.docx"
