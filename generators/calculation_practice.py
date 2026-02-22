"""
Calculation practice worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children practise arithmetic calculations in a clean 2-column grid layout.
"""

import io
import math
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_section_header,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_no_spacing,
    set_cell_shading,
    set_cell_borders,
    set_cell_padding,
    set_table_full_width,
    remove_table_borders,
)
from generators.styles import COLOURS, DIFF_LEVELS, FONT_NAME, THEMES
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_calculation_practice_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a calculation practice worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - sections: list of {title, instructions, calculations: [{question, answer, working_hint}]}
            - challenge: {title, instructions, lines: int} (optional)
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space
        show_answers: Whether to display answers (teacher answer key mode)

    Returns:
        BytesIO buffer containing the .docx file
    """
    diff = DIFF_LEVELS[level]
    theme = THEMES[theme_key]

    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    add_title_area(
        doc,
        content['title'],
        'Show your working and write your answers!',
        theme_key,
        level,
        is_answer_key=show_answers,
    )

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Add each calculation section
    for section_number, section in enumerate(content['sections'], start=1):
        # Section header
        add_section_header(doc, section_number, section['title'], theme_key)

        # Section instruction paragraph
        if section.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_inst = doc.add_paragraph()
            set_no_spacing(p_inst)
            p_inst.paragraph_format.space_after = Pt(6)
            run_inst = p_inst.add_run(section['instructions'])
            set_run_font(
                run_inst,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Build the 2-column calculation grid
        calculations = section.get('calculations', [])
        num_rows = math.ceil(len(calculations) / 2)

        if num_rows > 0:
            table = doc.add_table(rows=num_rows, cols=2)
            set_table_full_width(table)
            remove_table_borders(table)

            # Cell background: use the theme body colour for a light tint
            cell_bg = theme['body']
            cell_border = theme['accent']

            for idx, calc in enumerate(calculations):
                row_idx = idx // 2
                col_idx = idx % 2
                cell = table.cell(row_idx, col_idx)

                set_cell_shading(cell, cell_bg)
                set_cell_borders(cell, cell_border, sz=6)
                set_cell_padding(cell, top=120, bottom=120, left=150, right=150)

                # Question text
                p_q = cell.paragraphs[0]
                set_no_spacing(p_q)
                p_q.paragraph_format.space_after = Pt(4)
                run_q = p_q.add_run(calc['question'])
                set_run_font(
                    run_q,
                    size=Pt(diff['font_size']),
                    bold=True,
                    colour=COLOURS['black'],
                )

                # Show answer in bold green if answer key mode
                if show_answers and calc.get('answer'):
                    p_ans = cell.add_paragraph()
                    set_no_spacing(p_ans)
                    p_ans.paragraph_format.space_before = Pt(4)
                    run_ans = p_ans.add_run(f'Answer: {calc["answer"]}')
                    set_run_font(
                        run_ans,
                        size=Pt(diff['font_size']),
                        bold=True,
                        colour=COLOURS['criteria_text'],
                    )

                # Show working hint for developing level
                if calc.get('working_hint') and not show_answers:
                    p_hint = cell.add_paragraph()
                    set_no_spacing(p_hint)
                    p_hint.paragraph_format.space_before = Pt(4)
                    run_hint = p_hint.add_run(calc['working_hint'])
                    set_run_font(
                        run_hint,
                        size=Pt(diff['font_size'] - 3),
                        italic=True,
                        colour=COLOURS['hint_text'],
                    )

            # Clear any leftover empty cells when odd number of calculations
            if len(calculations) % 2 == 1:
                empty_cell = table.cell(num_rows - 1, 1)
                empty_cell.paragraphs[0].clear()

    # 5. Add challenge section if present (skip for developing level)
    challenge = content.get('challenge')
    if challenge and level != 'developing':
        challenge_number = len(content['sections']) + 1
        add_section_header(doc, challenge_number, challenge['title'], theme_key)

        # Challenge instructions
        if challenge.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_chal = doc.add_paragraph()
            set_no_spacing(p_chal)
            p_chal.paragraph_format.space_after = Pt(6)
            run_chal = p_chal.add_run(challenge['instructions'])
            set_run_font(
                run_chal,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Blank writing lines for the challenge
        num_lines = challenge.get('lines', 3)
        for _ in range(num_lines):
            p_line = doc.add_paragraph()
            set_no_spacing(p_line)
            p_line.paragraph_format.space_before = Pt(6)
            p_line.paragraph_format.space_after = Pt(6)
            run_line = p_line.add_run('_' * 70)
            set_run_font(
                run_line,
                size=Pt(diff['font_size']),
                colour=COLOURS['hint_text'],
            )

    # 6. Add success criteria checklist
    add_success_criteria(doc, content['success_criteria'], theme_key, level)

    # 7. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 8. Add footer with differentiation level label
    add_footer(doc, level, 'Calculation Practice')

    # 9. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
