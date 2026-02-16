"""
Word bank activity worksheet generator.

Takes structured JSON data (from an LLM) and produces a themed,
differentiated Word document using python-docx via reusable components.
Children learn new vocabulary and practise using words in sentences.
"""

import io
from generators.components import (
    create_base_document,
    add_title_area,
    add_learning_objective,
    add_word_bank,
    add_instructions,
    add_section_header,
    add_cloze_paragraph,
    add_success_criteria,
    add_eal_glossary_space,
    add_footer,
    set_run_font,
    set_no_spacing,
)
from generators.styles import COLOURS, DIFF_LEVELS
from docx.shared import Pt


# Instruction text varies by differentiation level
_INSTRUCTIONS = {
    'developing': 'Choose a word from the word bank to complete each sentence. Match the colours!',
    'expected': 'Use the words from the word bank to complete the sentences below.',
    'greater_depth': (
        'Use the word bank for ideas, then complete the sentences. '
        'Can you think of your own words too?'
    ),
}


def generate_word_bank_worksheet(
    content: dict,
    theme_key: str = 'classic',
    level: str = 'expected',
    objective: str = '',
    extra_spacing: bool = False,
    eal_glossary: bool = False,
    show_answers: bool = False,
) -> io.BytesIO:
    """
    Generate a word bank activity worksheet as a Word document.

    Args:
        content: Structured data from LLM with keys:
            - title: str
            - categories: list of {word_type, label, words: [{word, definition}]}
            - activities: list of {title, instructions, sentences: [{pieces}]}
            - success_criteria: list of str
        theme_key: Visual theme key (e.g. 'space', 'ocean', 'classic')
        level: Differentiation level ('developing', 'expected', 'greater_depth')
        objective: Learning objective text
        extra_spacing: Whether to add extra spacing for accessibility
        eal_glossary: Whether to include EAL glossary space

    Returns:
        BytesIO buffer containing the .docx file
    """
    diff = DIFF_LEVELS[level]

    # 1. Create the base document with standard margins and font
    doc = create_base_document(extra_spacing=extra_spacing)

    # 2. Add the themed title area with name/date fields
    add_title_area(
        doc,
        content['title'],
        'Learn new words and use them in sentences!',
        theme_key,
        level,
        is_answer_key=show_answers,
    )

    # 3. Add learning objective if provided
    if objective:
        add_learning_objective(doc, objective, theme_key)

    # 4. Add the colour-coded word bank with categories
    add_word_bank(doc, content['categories'], level)

    # 5. Add instructions with colour key
    instruction_text = _INSTRUCTIONS.get(level, _INSTRUCTIONS['expected'])
    add_instructions(doc, instruction_text, level)

    # 6. Add each activity: section header, instruction text, and cloze sentences
    for activity_number, activity in enumerate(content['activities'], start=1):
        # Section header
        add_section_header(doc, activity_number, activity['title'], theme_key)

        # Activity instruction paragraph
        if activity.get('instructions'):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(2)

            p_inst = doc.add_paragraph()
            set_no_spacing(p_inst)
            p_inst.paragraph_format.space_after = Pt(4)
            run_inst = p_inst.add_run(activity['instructions'])
            set_run_font(
                run_inst,
                size=Pt(diff['font_size'] - 2),
                italic=True,
                colour=COLOURS['grey_text'],
            )

        # Each sentence uses the cloze paragraph component (same pieces format)
        for sentence in activity['sentences']:
            pieces = sentence.get('pieces', sentence) if isinstance(sentence, dict) else sentence
            add_cloze_paragraph(doc, pieces, level, show_answers=show_answers)

    # 7. Add success criteria checklist
    add_success_criteria(doc, content['success_criteria'], theme_key, level)

    # 8. Add EAL glossary space if requested
    if eal_glossary:
        add_eal_glossary_space(doc)

    # 9. Add footer with differentiation level label
    add_footer(doc, level, 'Word Bank Activity')

    # 10. Save to BytesIO buffer and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
