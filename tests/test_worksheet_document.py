"""The worksheet she can actually hand out.

Found on 2026-09-04 by counting: the plan page had **one** download button, for
the lesson plan, and imported none of the ten worksheet generators. So the
headline feature — a worksheet built from the lesson, checked against it, and
proving her criteria — produced something she could look at on screen and could
not get out of the app. A worksheet that cannot be printed is not a worksheet.

The risk in adding the download is not that it fails; it is that it renders
**something other than the sheet that was checked.** The whole coupling
guarantee is that every criterion's quote was found in the sheet, so if the
document is built from anything but that same content, the guarantee stops
covering the thing the child is handed — which is exactly the hole closed on
2026-09-03, where the reply and the printed page had quietly stopped being the
same object.

So the test that matters is `test_every_quote_that_was_checked_is_on_the_page`,
which takes the evidence claims the guard accepted and finds each one in the
rendered document.

None of these touch the network.
"""

import io
import pathlib
import re

import pytest
from docx import Document

from planning.worksheet import (
    CoupledWorksheet,
    _normalise,
    _says_the_same,
    validate_coupled_worksheet,
)
from planning.worksheet_document import (
    GENERATOR_FOR,
    build_worksheet_document,
    worksheet_filename,
)
from planning.worksheet_schema import WORKSHEET_SCHEMAS
from tests.fixtures import ALL_CONTENT
from tests.test_library import CRITERION_ONE, CRITERION_TWO, INSTRUCTION, a_lesson

OBJECTIVE = "Compare the appearance of different rocks using property words"


def a_sheet(worksheet_type):
    """A checked worksheet of each kind, from the fixtures the generators render."""
    content = dict(ALL_CONTENT[worksheet_type])
    content["objective"] = OBJECTIVE
    content["success_criteria"] = [CRITERION_ONE]
    return CoupledWorksheet(
        objective=OBJECTIVE,
        success_criteria=[CRITERION_ONE],
        evidence=[],
        content=content,
        worksheet_type=worksheet_type,
        lesson_number=1,
    )


def blocks_in(document, normalise=True):
    """Every printed block — paragraphs and table cells alike.

    Kept as separate blocks rather than joined into one string, for the same
    reason the coupling check keeps the sheet as separate tasks: a quote
    stitched across the join between two unrelated boxes is not on the page.
    """
    document = Document(io.BytesIO(document.getvalue()))
    said = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            said += [cell.text for cell in row.cells]
    blocks = [block for block in said if block.strip()]
    return [_normalise(block) for block in blocks] if normalise else blocks


def words_in(document):
    """The whole page as one run. Only for asking whether a phrase is absent."""
    return " ".join(blocks_in(document))


#: What the page adds inside a gap that the reply never had: the word-type
#: symbol, and the hint in brackets beside it.
_ADDED_BY_THE_PAGE = re.compile(r"\([^)]*\)|[^\w\s'’,.!?:;—–\-/&%£+=<>\"]")


def as_the_child_reads_it(block):
    """One printed block, with what the page adds to a gap taken back off.

    Measured on 2026-09-04, by rendering and reading it: the printed sheet is
    **not** character-identical to the reply it was checked against. Every gap
    is drawn with the word-type symbol and the hint inside it, so

        "The rocket landed on the ___ surface of the planet."

    reaches the page as

        "The rocket landed on the  ⭐ __________ (a describing word about the
         surface)  surface of the planet."

    Both of those things are correct — the symbol is the dual coding that
    replaced the colour scheme, and the hint is the support. But it means the
    two strings can never be compared literally, and nothing in the project
    knew that until this test was written.
    """
    return _normalise(_ADDED_BY_THE_PAGE.sub(" ", block))


def printed_somewhere(quote, document):
    """Is this quote on the page, as the coupling check means it?

    ⚠️ Uses **the guard's own matcher**, deliberately, rather than a looser one
    written for this test. What is adjusted is the page, not the standard: the
    symbol and the hint come off, and then the quote has to say the same thing
    as one block does, by exactly the definition the guard uses. Loosening the
    match instead would let a quote that is genuinely not on the page pass.
    """
    return any(
        _says_the_same(quote, as_the_child_reads_it(block))
        for block in blocks_in(document, normalise=False)
    )


class TestEveryKindCanBePrinted:
    @pytest.mark.parametrize("worksheet_type", sorted(WORKSHEET_SCHEMAS))
    def test_a_sheet_of_every_kind_renders(self, worksheet_type):
        """Ten kinds, ten generators. A kind with no route out of the app is a
        kind she cannot use."""
        document = build_worksheet_document(a_sheet(worksheet_type))
        assert isinstance(document, io.BytesIO)
        assert document.getbuffer().nbytes > 0

    @pytest.mark.parametrize("worksheet_type", sorted(WORKSHEET_SCHEMAS))
    def test_every_kind_has_a_generator(self, worksheet_type):
        assert worksheet_type in GENERATOR_FOR

    @pytest.mark.parametrize("worksheet_type", sorted(WORKSHEET_SCHEMAS))
    def test_an_answer_key_can_be_printed_too(self, worksheet_type):
        assert build_worksheet_document(
            a_sheet(worksheet_type), show_answers=True
        ).getbuffer().nbytes > 0


class TestItIsHerLessonsSheet:
    def test_the_objective_printed_is_the_lessons(self):
        """Not re-derived and not the generator's own idea of one."""
        printed = words_in(build_worksheet_document(a_sheet("cloze")))
        assert _normalise(OBJECTIVE) in printed

    def test_her_criterion_is_printed(self):
        printed = words_in(build_worksheet_document(a_sheet("cloze")))
        assert _normalise(CRITERION_ONE) in printed

    def test_a_kind_with_no_generator_says_so(self):
        with pytest.raises(KeyError):
            build_worksheet_document(a_sheet("cloze").__class__(
                objective=OBJECTIVE, success_criteria=[], evidence=[],
                content={}, worksheet_type="a_kind_that_does_not_exist",
            ))


class TestTheSheetShePrintsIsTheSheetThatWasChecked:
    """⚠️ The one that carries the coupling guarantee out to the paper.

    Every criterion was accepted because its quote was found in the worksheet
    reply. That is only worth anything if the reply and the printed page are
    the same object — the exact thing that was **not** true on 2026-09-03, when
    six claims quoted sections no generator prints and the check passed them.
    """

    def _checked_sheet(self):
        """Built on the fixture the cloze generator really renders.

        ⚠️ Deliberately **not** the payload the coupling tests use. That one
        stores its sentences under `sentences`, which is enough to search for a
        quote in and is not a sheet the generator can draw — measured here,
        where it raised `KeyError: 'paragraphs'`. A test of "what she prints"
        has to start from content that prints.
        """
        lesson = a_lesson(OBJECTIVE, 1)
        payload = dict(ALL_CONTENT["cloze"])
        payload["objective"] = OBJECTIVE
        payload["success_criteria"] = [CRITERION_ONE, CRITERION_TWO]
        payload["evidence"] = [
            {
                "criterion": CRITERION_ONE,
                "where": "THE BEGINNING",
                "quote": "The rocket landed on the ___ surface of the planet.",
                "pupil_writes": "The describing word that fits",
            },
            {
                "criterion": CRITERION_TWO,
                "where": "THE BEGINNING",
                "quote": "Aisha scrambled down the ladder.",
                "pupil_writes": "The doing word that fits",
            },
        ]
        return validate_coupled_worksheet(
            payload, lesson=lesson, worksheet_type="cloze"
        )

    def test_every_quote_that_was_checked_is_on_the_page(self):
        """Each accepted quote, found in what she prints.

        The child's copy and the answer key are both looked in, and
        deliberately: a sentence with a gap is checked in either the form with
        the answers in place or the form with the gaps showing, so a quote may
        legitimately be either. What is not allowed is a quote that is in
        neither, because that is a criterion evidenced by nothing printed.
        """
        sheet = self._checked_sheet()
        pupils = build_worksheet_document(sheet)
        answers = build_worksheet_document(sheet, show_answers=True)

        for claim in sheet.evidence:
            assert printed_somewhere(claim.quote, pupils) or printed_somewhere(
                claim.quote, answers
            ), (
                f"{claim.criterion!r} was accepted on a quote that is not "
                f"printed on the sheet she hands out: {claim.quote!r}"
            )

    def test_it_would_notice_a_quote_that_is_not_printed(self):
        """The control. Without it the test above could pass on an empty page."""
        sheet = self._checked_sheet()
        printed = words_in(build_worksheet_document(sheet))
        assert _normalise("Colour in the rocks using your favourite colours.") not in printed


class TestTheTwoGeneratorMapsAgree:
    """⚠️ The map here is a second copy of the one in `app.py`.

    A deliberate, tested exception to this project's one-definition rule.
    `app.py` is a Streamlit script: importing it to borrow the map would run
    the whole worksheet flow, and there is a recorded decision that it stays
    structurally untouched because it works and has almost no coverage.

    So they are pinned to each other by reading `app.py` as text rather than
    running it — the same trick `RENDERED_KEYS` uses against the generators.
    Add an eleventh kind to one and not the other and this fails, rather than a
    teacher meeting a worksheet she cannot download.
    """

    def _map_declared_in_app(self):
        source = pathlib.Path(__file__).resolve().parent.parent / "app.py"
        text = source.read_text(encoding="utf-8")
        block = text.split("GENERATOR_MAP = {", 1)[1].split("}", 1)[0]
        return set(re.findall(r'"([a-z_]+)":', block))

    def test_app_declares_a_generator_map_at_all(self):
        """The control. If the parse silently found nothing, the test below
        would pass by comparing two empty sets."""
        assert len(self._map_declared_in_app()) == 10

    def test_the_same_kinds_are_in_both(self):
        assert self._map_declared_in_app() == set(GENERATOR_FOR)


class TestWhatLandsInHerDownloads:
    def test_the_file_is_named_after_the_unit_and_the_lesson(self):
        name = worksheet_filename(a_sheet("cloze"), "Rocks and Soils")
        assert name.endswith(".docx")
        assert "Rocks and Soils" in name
        assert "1" in name

    def test_the_answer_key_is_told_apart_from_the_childs_copy(self):
        sheet = a_sheet("cloze")
        assert worksheet_filename(sheet, "Rocks") != worksheet_filename(
            sheet, "Rocks", answers=True
        )
        assert "answers" in worksheet_filename(sheet, "Rocks", answers=True).lower()

    def test_a_title_she_typed_cannot_become_a_folder(self):
        """Same law as the lesson plan: the unit title is hers to type, and
        "Rocks and Soils: term 1/2" is a plausible thing to type."""
        name = worksheet_filename(a_sheet("cloze"), "Rocks/Soils: term 1")
        assert "/" not in name
        assert ":" not in name

    def test_a_unit_with_no_title_still_gets_a_usable_name(self):
        assert worksheet_filename(a_sheet("cloze"), "").endswith(".docx")
