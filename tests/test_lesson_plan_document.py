"""The lesson plan as a Word document she can actually edit.

Two decisions govern this file, both taken on 2026-09-01 and neither reopened.

**It has to be editable, not just printable.** Every box is a paragraph with a
border, never a one-cell table, and the formatting lives in named Word styles
rather than on each run. `prototype/make_editable_demo.py` builds the two
versions side by side: in the table version, pressing Enter inside a box fights
the table, a long answer stretches a cell instead of reflowing, and select-all
and change the font does nothing. She is a teacher with a document open the
night before, not a designer.

**Arial, black and blue.** The six-colour scheme is retired — the symbols and
labels that carried the meaning stay, so dual-coding survives and now also
survives a mono photocopier and a colour-blind child, which the colours never
did. Comic Sans and the joined font are both out; children still decoding, and
SEND children in particular, cannot read the joined one.

The test that matters most is neither of those. It is
`test_every_word_of_the_lesson_reaches_the_document`: the document must carry
everything the lesson holds. A generator that silently drops the misconceptions
produces a plausible, professional-looking plan with a hole in it, which is the
same failure as the lost lesson one layer down and just as invisible. That test
walks the lesson rather than naming fields, so a field added later and never
rendered fails it without anyone remembering to come back here.

None of these touch the network.
"""

import io
from dataclasses import fields, is_dataclass

import pytest
from docx import Document
from docx.oxml.ns import qn

from generators.lesson_plan import (
    ALLOWED_COLOURS,
    LESSON_PLAN_FONT,
    STYLE_PREFIX,
    build_lesson_plan,
    lesson_plan_bytes,
    lesson_plan_filename,
)
from planning.lesson import validate_lesson
from planning.worksheet import CoupledWorksheet, EvidenceClaim

OBJECTIVE = "Group rocks by their physical properties"
CRITERION_ONE = "I can describe two rocks using property words."
CRITERION_TWO = "I can put rocks into groups and say what the group has in common."


def step(minutes=20, **overrides):
    base = {
        "name": "Modelling",
        "minutes": minutes,
        "on_the_board": "Two rocks, and the words hard / soft / rough / smooth",
        "teacher_says": "Watch me. I am going to pick one word for this rock.",
        "questions": [
            {
                "ask": "Which word fits this rock?",
                "expect": "rough — you can feel the grains",
            }
        ],
        "children_do": "Talk to their partner and agree one word for the second rock",
        "watch_for": [
            {
                "wrong": "Children say 'nice' instead of a property",
                "respond": "Point back at the word bank and ask for a property word",
            }
        ],
        "adults": "TA sits with the four children on the back table",
        "builds_on_step": "The hook gave them the words; this shows how to use one",
    }
    base.update(overrides)
    return base


def a_lesson(**overrides):
    payload = {
        "objective": OBJECTIVE,
        "success_criteria": [
            {"criterion": CRITERION_ONE, "evidence": "completed comparison table"},
            {"criterion": CRITERION_TWO, "evidence": "sorted rocks stuck in with a heading"},
        ],
        "vocabulary": {
            "everyone": ["hard", "soft", "rough", "smooth"],
            "expected": ["grainy", "layered", "absorbent"],
            "stretch": ["permeable", "impermeable"],
            "guidance": "Teach the everyday word, then name the technical one beside it.",
        },
        "steps": [
            step(15, name="Hook"),
            step(20),
            step(20, name="Practice"),
            step(5, name="Plenary"),
        ],
        "misconceptions": [
            {
                "misconception": "Heavier means harder",
                "why": "Both are things you feel in your hands",
                "address": "Put chalk beside granite and let them scratch both",
            }
        ],
        "assessment": {
            "look_for": "A group heading that names a property rather than a colour",
            "not_yet_example": "Three rocks grouped under 'nice ones'",
        },
        "adaptations": {
            "eal": "Word bank with a photograph beside each property word",
            "send": "Pre-sorted pair of very different rocks to start from",
            "stretch": "Offer permeable and ask them to test it with a pipette",
        },
        "resources": [{"item": "Rock samples", "quantity": "6 sets of 4"}],
        "next_lesson": "Testing hardness by scratching",
    }
    payload.update(overrides)
    lesson = validate_lesson(
        payload, expected_objective=payload["objective"], lesson_minutes=60
    )
    return _with(lesson, **{k: v for k, v in overrides.items() if not isinstance(v, (dict, list))})


def _with(lesson, number=3, builds_on=2, builds_on_reason="", **_ignored):
    """The spine's fields, which `validate_lesson` does not set."""
    from planning.lesson import Lesson

    return Lesson(
        objective=lesson.objective,
        success_criteria=lesson.success_criteria,
        vocabulary=lesson.vocabulary,
        steps=lesson.steps,
        misconceptions=lesson.misconceptions,
        assessment=lesson.assessment,
        adaptations=lesson.adaptations,
        resources=lesson.resources,
        number=number,
        builds_on=builds_on,
        builds_on_reason=builds_on_reason
        or "Grouping has to be secure before a choice can be justified",
        next_lesson=lesson.next_lesson,
    )


def a_worksheet():
    return CoupledWorksheet(
        objective=OBJECTIVE,
        success_criteria=list(a_lesson().success_criteria),
        evidence=[
            EvidenceClaim(
                criterion=CRITERION_ONE,
                where="Part 1 — Describing",
                quote="Write two property words for each rock in the table.",
                pupil_writes="Two property words in each row",
            ),
            EvidenceClaim(
                criterion=CRITERION_TWO,
                where="Part 2 — Grouping",
                quote="Sort the six rocks into groups and write a heading for each group.",
                pupil_writes="A heading naming the property each group shares",
            ),
        ],
        content={"title": "Rock Detectives"},
        worksheet_type="cloze",
        lesson_number=3,
    )


def build(**overrides):
    defaults = dict(
        lesson=a_lesson(),
        unit_title="Rocks and Soils",
        subject="Science",
        year_group="Year 3",
        lesson_minutes=60,
        lesson_count=4,
        anchor="Boost — Rocks, unit 2",
        outcome="Children group rocks by their properties and say which rock they "
        "would choose for a job, and why.",
    )
    return build_lesson_plan(**{**defaults, **overrides})


def text_of(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def is_boxed(paragraph):
    """Does this paragraph carry a border of its own?

    Consecutive paragraphs with identical borders are merged by Word into one
    visual box, so "inside the box" is a property of each paragraph rather
    than of a container.
    """
    properties = paragraph._p.find(qn("w:pPr"))
    return properties is not None and properties.find(qn("w:pBdr")) is not None


# ── It has to be editable ────────────────────────────────────────────────────


class TestSheIsGoingToEditThis:
    """The decision the whole file turns on. See the prototype for the proof."""

    def test_the_document_has_no_tables(self):
        """A one-cell table used as a box is what makes the current worksheets
        a fight to edit. Real tables are for genuinely tabular content, and a
        lesson plan has none."""
        assert build().tables == []

    def test_every_paragraph_that_says_anything_carries_a_named_style(self):
        doc = build()
        unstyled = [
            p.text for p in doc.paragraphs
            if p.text.strip() and not p.style.name.startswith(STYLE_PREFIX)
        ]
        assert unstyled == [], (
            "Formatting on the run instead of in a style is what stops "
            "select-all-and-change-the-font from working."
        )

    def test_the_named_styles_are_defined_in_the_document(self):
        """A style referenced but not defined is formatting that vanishes the
        moment she opens the Styles pane."""
        doc = build()
        defined = {s.name for s in doc.styles}
        used = {p.style.name for p in doc.paragraphs if p.text.strip()}
        assert used <= defined

    def test_the_objective_style_carries_the_objective_and_nothing_else(self):
        """The styles are the deliverable, not decoration.

        The point of naming them is that she can restyle every objective box
        in the document at once from the Styles pane. A style called Objective
        that also carries the misconceptions and the worksheet's criteria
        restyles those too, silently, and the feature turns into a trap.
        Found reading a real generated plan on 2026-09-03.
        """
        doc = build(worksheet=a_worksheet())
        carried = [
            p.text.strip() for p in doc.paragraphs
            if p.style.name == "CA Objective"
        ]
        assert carried == [OBJECTIVE]

    def test_the_criterion_style_is_used_only_inside_the_objective_box(self):
        """Same reason. The resources list is a checklist, not a criterion."""
        doc = build()
        loose = [
            p.text.strip() for p in doc.paragraphs
            if p.style.name == "CA Criterion" and not is_boxed(p)
        ]
        assert loose == []


# ── Arial, black and blue ────────────────────────────────────────────────────


class TestTypography:
    def _fonts(self, doc):
        """Every font a reader can actually see.

        Not every font defined in the file: python-docx's blank template
        carries built-in styles nobody uses — `macro` is Courier in an empty
        document — and holding those against this file would be a check that
        can only ever fail. What counts is what the body asks for, what the
        styles the paragraphs actually use ask for, and `Normal`, which is what
        she inherits the moment she presses Enter at the end.
        """
        found = set()
        for element in doc.element.body.iter():
            if element.tag.endswith("}rFonts"):
                found.update(v for v in element.attrib.values() if v)
        used = {p.style.name for p in doc.paragraphs} | {"Normal"}
        for style in doc.styles:
            name = getattr(getattr(style, "font", None), "name", None)
            if name and style.name in used:
                found.add(name)
        return found

    def test_every_font_in_the_document_is_arial(self):
        assert self._fonts(build()) <= {LESSON_PLAN_FONT}

    def test_comic_sans_is_gone(self):
        """It is the default in the worksheet generators and must not leak in
        through a shared helper."""
        assert "Comic Sans MS" not in self._fonts(build())

    def test_the_joined_font_is_never_used(self):
        """Named in the decision: children still decoding, and SEND children in
        particular, cannot read it."""
        assert "Lucida Handwriting" not in self._fonts(build())

    def test_only_the_agreed_colours_appear(self):
        """Black and blue. Anything else is a colour a mono photocopier turns
        into an indistinguishable grey."""
        doc = build()
        used = set()
        for element in doc.element.body.iter():
            for key, value in element.attrib.items():
                if key.endswith("}color") or key.endswith("}fill"):
                    if value and value.lower() not in ("auto", "none"):
                        used.add(value.upper())
        assert used <= ALLOWED_COLOURS, f"Unexpected colours: {used - ALLOWED_COLOURS}"

    def test_the_font_can_be_changed(self):
        """Asked for by name: exposed as a setting she can change herself."""
        assert self._fonts(build(font="Verdana")) <= {"Verdana"}


# ── It carries the lesson ────────────────────────────────────────────────────


def _strings_in(value):
    """Every piece of text the lesson holds, wherever it is nested."""
    if isinstance(value, str):
        if value.strip():
            yield value.strip()
    elif isinstance(value, dict):
        for item in value.values():
            yield from _strings_in(item)
    elif isinstance(value, (list, tuple)):
        for item in value:
            yield from _strings_in(item)
    elif is_dataclass(value):
        for field in fields(value):
            yield from _strings_in(getattr(value, field.name))


class TestNothingIsQuietlyLeftOut:
    def test_every_word_of_the_lesson_reaches_the_document(self):
        """The one that would catch a hole nobody notices.

        A plan missing its misconceptions still looks like a plan — properly
        typeset, the right length, and wrong in a way she would only find in
        the room. This walks the lesson rather than naming its fields, so a
        field added later and never rendered fails here without anyone
        remembering to come back and add a test for it.

        It is deliberately a test and not a check in the generator. A guard
        that refused to produce the document over a rendering bug would take
        the plan away from her the night before she teaches it, which is worse
        than a plan that is slightly short.
        """
        lesson = a_lesson()
        rendered = text_of(build(lesson=lesson))
        missing = [
            text for text in dict.fromkeys(_strings_in(lesson))
            if text not in rendered
        ]
        assert missing == [], f"The document leaves out: {missing}"

    def test_the_objective_is_hers_word_for_word(self):
        """Three handovers end here. The plan, the sheet and the child's book
        have to carry the same sentence."""
        assert OBJECTIVE in [p.text.strip() for p in build().paragraphs]

    def test_the_criteria_are_printed_in_her_order(self):
        rendered = text_of(build())
        assert rendered.index(CRITERION_ONE) < rendered.index(CRITERION_TWO)

    def test_the_three_vocabulary_bands_are_named_separately(self):
        """One list is precisely what she said does not work for this class."""
        rendered = text_of(build()).lower()
        for band in ("everyone", "expected", "stretch"):
            assert band in rendered

    def test_the_guidance_on_using_the_bands_says_that_is_what_it_is(self):
        """Found reading a real generated plan on 2026-09-03: unlabelled, it
        sits under the stretch band and reads as a fourth list of stretch
        words rather than as how to use the three."""
        assert "How to use them:" in text_of(build())

    def test_every_step_says_how_long_it_takes(self):
        rendered = text_of(build())
        for minutes in (15, 20, 5):
            assert f"{minutes} min" in rendered

    def test_the_steps_are_in_the_order_they_happen(self):
        rendered = text_of(build())
        assert rendered.index("Hook") < rendered.index("Practice") < rendered.index("Plenary")

    def test_it_is_labelled_a_draft_and_never_verified(self):
        rendered = text_of(build()).lower()
        assert "check before teaching" in rendered
        assert "verified" not in rendered


# ── Where the lesson sits ────────────────────────────────────────────────────


class TestTheUnitAroundIt:
    def test_it_says_which_unit_and_where_in_it(self):
        rendered = text_of(build())
        assert "Rocks and Soils" in rendered
        assert "3 of 4" in rendered

    def test_it_says_where_the_lesson_came_from(self):
        """Her evidence to a subject leader that the scheme was followed."""
        assert "Boost — Rocks, unit 2" in text_of(build())

    def test_it_says_what_this_lesson_needs_and_why(self):
        rendered = text_of(build())
        assert "Grouping has to be secure before a choice can be justified" in rendered

    def test_the_first_lesson_does_not_claim_to_build_on_anything(self):
        lesson = _with(a_lesson(), number=1, builds_on=None, builds_on_reason="")
        assert "builds on" not in text_of(build(lesson=lesson)).lower()

    def test_the_unit_outcome_is_carried(self):
        assert "would choose for a job" in text_of(build())


# ── The worksheet it goes with ───────────────────────────────────────────────


class TestTheWorksheetItGoesWith:
    """The headline feature, said on the page she prints.

    The worksheet now carries which task evidences which criterion, so the
    plan can say it too — and a plan that says which sheet proves which
    criterion is what makes the pair defensible to a moderator.
    """

    def test_it_says_which_task_evidences_which_criterion(self):
        rendered = text_of(build(worksheet=a_worksheet()))
        assert "Write two property words for each rock in the table." in rendered
        assert "Sort the six rocks into groups and write a heading for each group." in rendered

    def test_it_names_the_part_of_the_sheet_and_what_the_child_writes(self):
        rendered = text_of(build(worksheet=a_worksheet()))
        assert "Part 1 — Describing" in rendered
        assert "Two property words in each row" in rendered

    def test_without_a_worksheet_it_says_nothing_about_one(self):
        """A plan that mentions a sheet she has not made is a plan that lies."""
        assert "worksheet" not in text_of(build()).lower()

    def test_a_worksheet_for_a_different_lesson_is_refused(self):
        """The one guard here, and it is not about formatting: a plan that
        prints another lesson's sheet as this lesson's evidence is wrong in a
        way she would have to teach to discover."""
        sheet = a_worksheet()
        other = CoupledWorksheet(
            objective="Something else entirely",
            success_criteria=sheet.success_criteria,
            evidence=sheet.evidence,
            content=sheet.content,
            worksheet_type=sheet.worksheet_type,
            lesson_number=sheet.lesson_number,
        )
        with pytest.raises(ValueError, match="objective"):
            build(worksheet=other)


# ── It has to open ───────────────────────────────────────────────────────────


class TestTheFileSheGets:
    """What lands in her Downloads folder.

    Checked here rather than on the screen because AppTest cannot see a
    download button's file name — the file is served over a URL and the
    element carries only a deferred id.
    """

    def test_it_is_named_after_the_unit_and_the_lesson(self):
        """A folder of files called `download.docx` is a folder of nothing."""
        assert lesson_plan_filename(a_lesson(), "Rocks and Soils") == (
            "Rocks and Soils - lesson 3 - plan.docx"
        )

    def test_a_unit_with_no_title_still_names_the_lesson(self):
        assert lesson_plan_filename(a_lesson(), "") == "Lesson 3 - plan.docx"

    def test_a_slash_in_the_unit_title_does_not_become_a_folder(self):
        """She types the unit title. "Rocks/Soils" is a plausible thing to
        type and a path separator on every operating system she uses."""
        name = lesson_plan_filename(a_lesson(), "Rocks/Soils: term 1")
        assert "/" not in name and ":" not in name
        assert name.endswith(".docx")


class TestItOpens:
    def test_the_bytes_are_a_document_word_can_open(self):
        data = lesson_plan_bytes(
            lesson=a_lesson(),
            unit_title="Rocks and Soils",
            subject="Science",
            year_group="Year 3",
        )
        reopened = Document(io.BytesIO(data))
        assert OBJECTIVE in text_of(reopened)
        assert reopened.tables == []
