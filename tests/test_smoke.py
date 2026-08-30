"""
End-to-end smoke tests: every worksheet type, every differentiation level,
every theme -- built into a real .docx and opened again to prove it is valid.

No API key and no network. The Claude call is the only part of the pipeline
these do not cover; ``test_sdk_contract.py`` guards that boundary separately.
"""

import ast
import io
import pathlib

import docx
import pytest

from generators.calculation_practice import generate_calculation_practice_worksheet
from generators.cloze import generate_cloze_worksheet
from generators.fraction_practice import generate_fraction_practice_worksheet
from generators.investigation import generate_investigation_worksheet
from generators.matching import generate_matching_worksheet
from generators.problem_solving import generate_problem_solving_worksheet
from generators.reading_comprehension import generate_reading_comprehension_worksheet
from generators.sentence_builder import generate_sentence_builder_worksheet
from generators.styles import DIFF_LEVELS, THEMES, YEAR_AGES
from generators.times_tables import generate_times_tables_worksheet
from generators.word_bank import generate_word_bank_worksheet
from llm.prompts import get_prompt, list_worksheet_types
from tests.fixtures import ALL_CONTENT

REPO_ROOT = pathlib.Path(__file__).resolve().parent.parent

GENERATORS = {
    "cloze": generate_cloze_worksheet,
    "word_bank": generate_word_bank_worksheet,
    "matching": generate_matching_worksheet,
    "sentence_builder": generate_sentence_builder_worksheet,
    "reading_comprehension": generate_reading_comprehension_worksheet,
    "problem_solving": generate_problem_solving_worksheet,
    "calculation_practice": generate_calculation_practice_worksheet,
    "fraction_practice": generate_fraction_practice_worksheet,
    "times_tables": generate_times_tables_worksheet,
    "investigation": generate_investigation_worksheet,
}

LEVELS = list(DIFF_LEVELS)
THEME_KEYS = list(THEMES)


def _document_text(buffer):
    """Open a generated buffer as a Word document and return all of its text."""
    buffer.seek(0)
    document = docx.Document(buffer)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# --------------------------------------------------------------------------
# The map in app.py must stay in step with the generators and the prompts.
# Parsed from source rather than imported, because importing app.py would run
# the whole Streamlit script.
# --------------------------------------------------------------------------


def _app_generator_map_keys():
    tree = ast.parse((REPO_ROOT / "app.py").read_text())
    for node in ast.walk(tree):
        if isinstance(node, ast.Assign):
            for target in node.targets:
                if isinstance(target, ast.Name) and target.id == "GENERATOR_MAP":
                    return {k.value for k in node.value.keys}
    raise AssertionError("GENERATOR_MAP not found in app.py")


def test_app_offers_exactly_the_worksheet_types_we_can_generate():
    assert _app_generator_map_keys() == set(GENERATORS)


def test_every_worksheet_type_has_a_prompt():
    """A type the app offers but has no prompt for is a dead menu entry."""
    missing = set(GENERATORS) - set(list_worksheet_types())
    assert not missing, f"Worksheet types with no prompt: {sorted(missing)}"


def test_every_worksheet_type_has_test_content():
    """Stops a new worksheet type shipping with no smoke coverage."""
    missing = set(GENERATORS) - set(ALL_CONTENT)
    assert not missing, (
        f"No test content for: {sorted(missing)}. Add a fixture in "
        "tests/fixtures.py so this type is covered."
    )


# --------------------------------------------------------------------------
# Document generation -- the bulk of the app.
# --------------------------------------------------------------------------


@pytest.mark.parametrize("ws_type", sorted(GENERATORS))
@pytest.mark.parametrize("level", LEVELS)
def test_generates_a_valid_document_for_every_type_and_level(ws_type, level):
    buffer = GENERATORS[ws_type](
        content=ALL_CONTENT[ws_type],
        theme_key="space",
        level=level,
        objective="Pupils can use precise vocabulary to describe a setting.",
        extra_spacing=False,
        eal_glossary=False,
        show_answers=False,
    )

    assert isinstance(buffer, io.BytesIO)
    text = _document_text(buffer)
    assert text.strip(), f"{ws_type} at {level} produced an empty document"
    assert ALL_CONTENT[ws_type]["title"] in text, (
        f"{ws_type} at {level} did not render its own title"
    )


@pytest.mark.parametrize("ws_type", sorted(GENERATORS))
def test_answer_key_renders_and_differs_from_the_pupil_sheet(ws_type):
    """The teacher edition must actually contain something the pupil sheet does not."""
    common = dict(
        content=ALL_CONTENT[ws_type],
        theme_key="classic",
        level="expected",
        objective="",
        extra_spacing=False,
        eal_glossary=False,
    )
    pupil = _document_text(GENERATORS[ws_type](show_answers=False, **common))
    teacher = _document_text(GENERATORS[ws_type](show_answers=True, **common))

    assert teacher.strip(), f"{ws_type} answer key was empty"
    assert teacher != pupil, (
        f"{ws_type} answer key is identical to the pupil sheet -- show_answers "
        "is not doing anything."
    )


@pytest.mark.parametrize("theme_key", THEME_KEYS)
def test_every_theme_renders(theme_key):
    """A broken theme should not wait for a teacher to discover it."""
    buffer = GENERATORS["cloze"](
        content=ALL_CONTENT["cloze"],
        theme_key=theme_key,
        level="expected",
        objective="",
        extra_spacing=False,
        eal_glossary=False,
        show_answers=False,
    )
    assert _document_text(buffer).strip()


@pytest.mark.parametrize("ws_type", sorted(GENERATORS))
def test_accessibility_options_render(ws_type):
    """Extra spacing and the EAL glossary are the two accessibility toggles."""
    buffer = GENERATORS[ws_type](
        content=ALL_CONTENT[ws_type],
        theme_key="ocean",
        level="developing",
        objective="An objective.",
        extra_spacing=True,
        eal_glossary=True,
        show_answers=False,
    )
    assert _document_text(buffer).strip()


# --------------------------------------------------------------------------
# Prompt building.
# --------------------------------------------------------------------------


@pytest.mark.parametrize("ws_type", sorted(GENERATORS))
@pytest.mark.parametrize("year_group", sorted(YEAR_AGES))
def test_prompt_builds_for_every_type_and_year(ws_type, year_group):
    prompt = get_prompt(
        ws_type,
        year_group=year_group,
        topic="Writing - Myths and Legends",
        objective="Plan, draft and write a narrative opening.",
        age_range=YEAR_AGES[year_group],
        theme_name="Space Explorer",
        theme_icon="rocket",
        level="expected",
        subject="English",
    )
    assert isinstance(prompt, str)
    assert prompt.strip(), f"{ws_type} produced an empty prompt for {year_group}"
    assert year_group in prompt, f"{ws_type} prompt did not mention {year_group}"


def test_unknown_worksheet_type_is_rejected_clearly():
    with pytest.raises(ValueError, match="Unknown worksheet type"):
        get_prompt("interpretive_dance", year_group="Year 3")


# --------------------------------------------------------------------------
# Dependency pinning -- the root cause of the outage this suite was written for.
# --------------------------------------------------------------------------


def test_dependencies_have_upper_bounds():
    """An unbounded pin let a major SDK release break the live app silently.

    Every runtime dependency must be capped below the next major version so a
    rebuild cannot pull in breaking changes without a deliberate edit here.
    """
    requirements = (REPO_ROOT / "requirements.txt").read_text().splitlines()
    unbounded = [
        line.strip()
        for line in requirements
        if line.strip() and not line.strip().startswith("#") and "<" not in line
    ]
    assert not unbounded, (
        f"These dependencies have no upper bound: {unbounded}. A new major "
        "version can install itself on the next rebuild and break the app with "
        "no code change. Add a '<N.0.0' cap to each."
    )
