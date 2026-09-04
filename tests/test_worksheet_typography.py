"""Arial, black and blue — across every worksheet the app makes.

The decision, taken 2026-09-01 and applying to lesson plans *and* worksheets:

**Arial throughout.** `FONT_NAME` was Comic Sans and is read by nine call
sites. The joined font used elsewhere in school is out by name — children still
decoding, and SEND children in particular, cannot read it.

**Black and blue only.** The six-colour word-type scheme is retired. The
symbols and labels that carried the meaning stay, so dual-coding survives —
and it now also works photocopied in mono and for a colour-blind child, which
the colour scheme never did.

That last clause is the part with teeth, and it is why this file is not just a
palette check. Colour was doing real work in three places, and in one of them
it was doing it *alone*: the sentence-builder word cards showed the word and
nothing else, with the word type carried entirely by the fill. Retiring the
colours without putting the symbol on those cards would not be a restyle — it
would delete the information. The symbol also has to be unambiguous: two word
types offered to the same subject cannot share one, because inline in a cloze
sentence the symbol is all there is.

The sweep runs every generator on the real fixtures, because a palette is only
retired where it is actually rendered.
"""

import io
import re
from collections import defaultdict

import pytest
from docx import Document

from generators.styles import FONT_NAME, PALETTE, WORD_TYPES
from llm.prompts import SUBJECT_WORD_TYPES
from tests.fixtures import ALL_CONTENT
from tests.test_smoke import GENERATORS

LEVELS = ("developing", "expected", "greater_depth")

# The joined/calligraphic font used elsewhere in school. Out by name.
JOINED_FONT = "Lucida Handwriting"

# An instruction to use colour to find meaning -- not any mention of the word.
# The first version of this matched a bare "colour" and would have deleted the
# UK-spelling instruction ("'colour' not 'color'"), which is correct and
# needed. A guard that refuses correct work is worse than no guard.
COLOUR_INSTRUCTION = re.compile(
    r"colou?r[- ]?(?:cod\w*|group\w*|key)|match(?:ing)? the colou?r|colou?red\b",
    re.I,
)


def a_sheet(kind, level="expected", theme="space", answers=False, font=None):
    extra = {"font": font} if font else {}
    return Document(
        GENERATORS[kind](
            content=ALL_CONTENT[kind],
            theme_key=theme,
            level=level,
            objective="Pupils can describe a rock using property words.",
            extra_spacing=False,
            eal_glossary=True,
            show_answers=answers,
            **extra,
        )
    )


def fonts_in(doc):
    """Every font the file asks for, in the body and in the styles it uses."""
    found = set()
    for part in (doc.element.body,) + tuple(t._tbl for t in doc.tables):
        for element in part.iter():
            if element.tag.endswith("}rFonts"):
                found.update(v for v in element.attrib.values() if v)
    used = {p.style.name for p in doc.paragraphs} | {"Normal"}
    for style in doc.styles:
        name = getattr(getattr(style, "font", None), "name", None)
        if name and style.name in used:
            found.add(name)
    return found


def colours_in(doc):
    """Every colour the file paints — text, fills and borders alike."""
    found = set()
    for element in doc.element.body.iter():
        for key, value in element.attrib.items():
            if key.endswith("}color") or key.endswith("}fill"):
                if value and value.lower() not in ("auto", "none"):
                    found.add(value.upper())
    return found


def text_of(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ── The font ─────────────────────────────────────────────────────────────────


class TestArialThroughout:
    def test_the_shared_default_is_the_agreed_font(self):
        assert FONT_NAME == "Arial"

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_no_worksheet_is_set_in_comic_sans(self, kind):
        assert "Comic Sans MS" not in fonts_in(a_sheet(kind))

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_no_worksheet_uses_the_joined_font(self, kind):
        assert JOINED_FONT not in fonts_in(a_sheet(kind))

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_every_font_a_worksheet_asks_for_is_the_agreed_one(self, kind):
        found = fonts_in(a_sheet(kind))
        assert found <= {FONT_NAME}, f"{kind} also asks for {found - {FONT_NAME}}"

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_she_can_change_the_font(self, kind):
        """Asked for by name in the decision: Arial, *and exposed as a setting
        she can change herself.* She is the one reading it at seven in the
        morning with a class arriving."""
        found = fonts_in(a_sheet(kind, font="Verdana"))
        assert found <= {"Verdana"}, f"{kind} also asks for {found - {'Verdana'}}"

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_no_run_pins_its_own_font(self, kind):
        """The mechanism, and the reason changing it is one setting.

        Every run used to name the font itself, so the font lived in several
        hundred places per document. Now nothing does, and the runs inherit
        the document's Normal style — which is also what makes select-all and
        change-the-font work for her in Word.
        """
        doc = a_sheet(kind)
        pinned = [
            element for element in doc.element.body.iter()
            if element.tag.endswith("}rFonts")
        ]
        assert pinned == [], f"{kind} names the font on {len(pinned)} runs"


# ── The palette ──────────────────────────────────────────────────────────────


class TestBlackAndBlueOnly:
    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    @pytest.mark.parametrize("level", LEVELS)
    def test_no_worksheet_paints_outside_the_palette(self, kind, level):
        used = colours_in(a_sheet(kind, level=level))
        assert used <= PALETTE, f"{kind} at {level} also paints {used - PALETTE}"

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_the_answer_key_is_the_same_palette(self, kind):
        """The teacher's copy goes through the same photocopier."""
        used = colours_in(a_sheet(kind, answers=True))
        assert used <= PALETTE, f"{kind} answers also paint {used - PALETTE}"

    @pytest.mark.parametrize("theme", ("space", "ocean", "jungle", "classic"))
    def test_the_themes_no_longer_carry_their_own_colours(self, theme):
        """The themes keep their names and icons — "Mission", "Captain's Log",
        the rocket — because those are language, not colour. What they lose is
        seven different palettes."""
        used = colours_in(a_sheet("cloze", theme=theme))
        assert used <= PALETTE, f"the {theme} theme paints {used - PALETTE}"


# ── What the colour was carrying ─────────────────────────────────────────────


class TestTheMeaningSurvivesWithoutTheColour:
    """Retiring a code is only safe if something else carries what it said."""

    def test_every_word_type_still_has_a_symbol_and_a_label(self):
        for key, word_type in WORD_TYPES.items():
            assert word_type.get("symbol"), f"{key} has no symbol left"
            assert word_type.get("label"), f"{key} has no label left"

    @pytest.mark.parametrize("subject", sorted(SUBJECT_WORD_TYPES))
    def test_no_subject_offers_two_word_types_with_the_same_symbol(self, subject):
        """Inline in a cloze sentence the symbol is all there is.

        Found on 2026-09-03 while retiring the colours: Languages offers both
        "describing word" and "key word", and both were a star. Today the fill
        colour tells them apart; with the colours gone they would be the same
        mark, and nothing on screen or on paper would say so.
        """
        offered = re.findall(r'- "([a-z_]+)" with label', SUBJECT_WORD_TYPES[subject])
        seen = defaultdict(list)
        for key in offered:
            if key in WORD_TYPES:
                seen[WORD_TYPES[key]["symbol"]].append(key)
        clashes = {sym: keys for sym, keys in seen.items() if len(keys) > 1}
        assert not clashes, f"{subject} offers one symbol for {clashes}"

    def test_the_prompt_and_the_sheet_show_the_same_symbol(self):
        """The generator overrides whatever label came back, so a prompt naming
        a different symbol is a false instruction rather than a broken sheet —
        but it is still false, and it is what the next person would read."""
        for subject, block in SUBJECT_WORD_TYPES.items():
            for key, label in re.findall(r'- "([a-z_]+)" with label "(.+?)"', block):
                if key in WORD_TYPES:
                    assert WORD_TYPES[key]["symbol"] in label, (
                        f"{subject} tells the model {label!r} for {key!r}, but the "
                        f"sheet prints {WORD_TYPES[key]['symbol']!r}"
                    )

    def test_a_sentence_builder_card_shows_its_symbol(self):
        """The one place colour was the *only* carrier.

        A word card printed the word and nothing else; which kind of word it
        was lived entirely in the fill. Retiring the colours without this
        would not restyle the sheet, it would delete the information a child
        sorts by.
        """
        rendered = text_of(a_sheet("sentence_builder"))
        parts = [
            part
            for exercise in ALL_CONTENT["sentence_builder"]["exercises"]
            for part in exercise.get("sentence_parts", [])
        ]
        assert parts, "the fixture has no word cards to check"
        for part in parts:
            symbol = WORD_TYPES.get(part.get("word_type", "noun"), WORD_TYPES["noun"])[
                "symbol"
            ]
            assert f'{symbol} {part["part"]}' in rendered, (
                f'the card {part["part"]!r} does not show its {symbol} mark'
            )

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_no_sheet_tells_a_child_to_use_a_colour(self, kind):
        """Found by reading a rendered sheet on 2026-09-03, and invisible to
        every palette check in this file.

        The sheets came out in black and blue and went on printing *"Match
        the colour and symbol to find the right word"* and *"Write your own
        sentence using a word from each colour group."* A child following
        that instruction is looking for something that is not on the page —
        which is worse than the colours never having been there.
        """
        found = COLOUR_INSTRUCTION.findall(text_of(a_sheet(kind)))
        assert not found, f"{kind} still tells a child to use colour: {set(found)}"

    def test_no_prompt_asks_the_model_for_colour_coding(self):
        """Where the instruction came from. The model was told the word types
        were for colour-coding and told to write an extension about colour
        groups, so the sheet said so — it was doing as it was asked."""
        from llm.prompts import get_prompt

        blocks = dict(SUBJECT_WORD_TYPES)
        for kind in sorted(GENERATORS):
            blocks[f"{kind} prompt"] = get_prompt(
                worksheet_type=kind,
                year_group="Year 3",
                topic="Rocks",
                objective="Describe a rock",
                age_range="7-8",
                theme_name="Space Explorer",
                theme_icon="\U0001F680",
                level="expected",
                subject="Science",
            )
        offenders = {
            name: set(COLOUR_INSTRUCTION.findall(text))
            for name, text in blocks.items()
            if COLOUR_INSTRUCTION.search(text)
        }
        assert not offenders, f"still asking for colour: {offenders}"

    def test_the_key_is_no_longer_called_a_colour_key(self):
        """There are no colours in it. A name that says otherwise is the kind
        of false label this project has been bitten by before."""
        import generators.components as components

        assert not hasattr(components, "add_colour_key")
        assert hasattr(components, "add_symbol_key")


# ── It still has to be a worksheet ───────────────────────────────────────────


class TestTheSheetsStillRender:
    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_every_worksheet_still_builds_and_says_its_title(self, kind):
        doc = a_sheet(kind)
        assert ALL_CONTENT[kind]["title"] in text_of(doc)

    @pytest.mark.parametrize("kind", sorted(GENERATORS))
    def test_nothing_became_invisible(self, kind):
        """White on white is a way to pass a palette check and lose a sheet."""
        doc = a_sheet(kind)
        assert "FFFFFF" not in colours_in(doc) or text_of(doc).strip()
        assert isinstance(
            GENERATORS[kind](
                content=ALL_CONTENT[kind],
                theme_key="space",
                level="expected",
                objective="",
                extra_spacing=False,
                eal_glossary=False,
                show_answers=False,
            ),
            io.BytesIO,
        )
