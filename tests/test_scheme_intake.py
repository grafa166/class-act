"""Reading the plan the teacher has to follow.

Boost's long- and medium-term plans exist and must be followed, but they are
too thin to teach from. So the tool takes them as a constraint: the coverage
and order are kept because she is accountable for them, and the teaching is
rebuilt.

She should not have to retype any of it. Three ways in -- paste the text,
upload the file the subject leader circulated, or photograph the printed page,
because a photocopy in a folder is the realistic case. PDFs and photographs go
to Claude as document and image blocks rather than through a text extractor:
a scanned or photographed plan has no text layer for a parser to find.

None of these tests touch the network. The one function that would --
`read_scheme_plan` -- takes its model call from a module-level name, so the
tests replace it and inspect exactly what would have been sent.
"""

import base64
import io

import pytest
from docx import Document

import planning.scheme_intake as scheme_intake
from planning.scheme_intake import (
    EXTRACTION_SYSTEM_PROMPT,
    MAX_UPLOAD_BYTES,
    SchemePlan,
    SchemePlanError,
    UnreadableUploadError,
    blocks_for_upload,
    build_extraction_prompt,
    read_scheme_plan,
    vague_coverage_items,
    validate_scheme_plan,
)

GOOD = {
    "unit_title": "Rocks and Soils",
    "coverage": [
        "types of rock",
        "properties",
        "fossils",
        "soil formation",
    ],
    "assessment": ["children know that rocks have different properties"],
    "activities": ["rock sorting", "fossil worksheet", "soil jar"],
}


class TestReadingUploads:
    def test_a_pdf_becomes_a_document_block(self):
        """Not text-extracted. A photocopied plan has no text layer."""
        blocks = blocks_for_upload("plan.pdf", b"%PDF-1.4 fake")
        assert blocks[0]["type"] == "document"
        assert blocks[0]["source"]["media_type"] == "application/pdf"

    def test_a_photograph_becomes_an_image_block(self):
        blocks = blocks_for_upload("plan.jpg", b"\xff\xd8\xff fake jpeg")
        assert blocks[0]["type"] == "image"
        assert blocks[0]["source"]["media_type"] == "image/jpeg"

    def test_png_is_recognised_separately_from_jpeg(self):
        blocks = blocks_for_upload("plan.png", b"\x89PNG fake")
        assert blocks[0]["source"]["media_type"] == "image/png"

    def test_base64_carries_no_newlines(self):
        """The API rejects wrapped base64, and b64encode is not the trap --
        textwrap-style wrapping added by a helper is."""
        blocks = blocks_for_upload("plan.pdf", b"x" * 500)
        assert "\n" not in blocks[0]["source"]["data"]

    def test_the_payload_round_trips(self):
        raw = b"some bytes that must survive"
        blocks = blocks_for_upload("plan.pdf", raw)
        assert base64.b64decode(blocks[0]["source"]["data"]) == raw

    def test_plain_text_is_passed_through_as_text(self):
        blocks = blocks_for_upload("plan.txt", "Rocks and Soils\nCoverage: ...".encode())
        assert blocks[0]["type"] == "text"
        assert "Rocks and Soils" in blocks[0]["text"]

    def test_an_unsupported_type_is_refused_by_name(self):
        with pytest.raises(UnreadableUploadError, match="xlsx"):
            blocks_for_upload("plan.xlsx", b"...")

    def test_an_oversized_upload_is_refused_before_it_is_encoded(self):
        with pytest.raises(UnreadableUploadError, match="too large"):
            blocks_for_upload("plan.pdf", b"x" * (MAX_UPLOAD_BYTES + 1))

    def test_an_empty_upload_is_refused(self):
        with pytest.raises(UnreadableUploadError):
            blocks_for_upload("plan.pdf", b"")

    def test_extension_matching_ignores_case(self):
        assert blocks_for_upload("PLAN.PDF", b"%PDF")[0]["type"] == "document"


class TestExtractionPrompt:
    def _prompt(self, **kw):
        return build_extraction_prompt(
            scheme="Boost", subject="Science", year_group="Year 3", **kw
        )

    def test_it_names_the_scheme_and_year(self):
        p = self._prompt()
        assert "Boost" in p and "Year 3" in p and "Science" in p

    def test_it_forbids_inventing_coverage(self):
        """The whole point is fidelity to what the school agreed to teach."""
        low = self._prompt().lower()
        assert "do not invent" in low or "only what" in low

    def test_it_asks_for_the_vague_items_to_be_flagged(self):
        assert "vague" in self._prompt().lower()

    def test_it_requests_json_only(self):
        assert "json" in self._prompt().lower()

    def test_flagging_a_line_must_not_remove_it_from_the_coverage(self):
        """Measured against the live API on 2026-09-02: told to 'flag a vague
        entry in the vague list', Claude *moved* the line there and left it out
        of the coverage -- so a lesson the school had planned disappeared with
        nothing on screen saying so."""
        low = self._prompt().lower()
        assert "move" in low and "both" in low


class TestValidation:
    def test_a_good_payload_becomes_a_scheme_plan(self):
        plan = validate_scheme_plan(GOOD)
        assert isinstance(plan, SchemePlan)
        assert plan.unit_title == "Rocks and Soils"
        assert len(plan.coverage) == 4

    def test_coverage_may_not_be_empty(self):
        """No coverage means nothing to be accountable to -- reject it."""
        with pytest.raises(SchemePlanError, match="coverage"):
            validate_scheme_plan({**GOOD, "coverage": []})

    def test_a_missing_unit_title_is_rejected(self):
        with pytest.raises(SchemePlanError):
            validate_scheme_plan({**GOOD, "unit_title": "   "})

    def test_coverage_must_be_strings_not_nested_objects(self):
        with pytest.raises(SchemePlanError):
            validate_scheme_plan({**GOOD, "coverage": [{"item": "rocks"}]})

    def test_a_non_dict_payload_is_rejected(self):
        with pytest.raises(SchemePlanError):
            validate_scheme_plan(["rocks"])

    def test_assessment_and_activities_are_optional(self):
        plan = validate_scheme_plan(
            {"unit_title": "Rocks", "coverage": ["types of rock"]}
        )
        assert plan.assessment == [] and plan.activities == []

    def test_blank_coverage_entries_are_dropped_not_kept(self):
        plan = validate_scheme_plan({**GOOD, "coverage": ["rocks", "  ", ""]})
        assert plan.coverage == ["rocks"]

    def test_dropping_blanks_cannot_empty_the_coverage_silently(self):
        with pytest.raises(SchemePlanError):
            validate_scheme_plan({**GOOD, "coverage": ["  ", ""]})


class TestVagueness:
    def test_a_statement_with_no_verb_for_the_child_is_flagged(self):
        """'Children know that rocks have different properties' is a statement,
        not an objective -- nothing in it says what a child would do."""
        flagged = vague_coverage_items(["children know that rocks have properties"])
        assert flagged

    def test_a_bare_topic_label_is_flagged(self):
        assert vague_coverage_items(["properties"])

    def test_a_teachable_objective_is_not_flagged(self):
        assert not vague_coverage_items(
            ["compare and group rocks by their physical properties"]
        )

    def test_flagging_explains_itself(self):
        """A flag the teacher cannot argue with is worse than no flag."""
        (item, reason), = vague_coverage_items(["properties"])
        assert item == "properties"
        assert reason.strip()

    def test_it_is_a_hint_not_a_verdict(self):
        """Documented behaviour: nothing here rejects a unit. The teacher
        decides -- this only surfaces what to look at."""
        assert isinstance(vague_coverage_items(["properties"]), list)


def _docx_bytes(build):
    """A real .docx in memory, built by `build(document)`."""
    document = Document()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestReadingAWordDocument:
    """The file the subject leader circulated is the likeliest upload of all.

    Unlike a PDF or a photograph, a .docx has a real text layer and the API has
    no block type for it, so this one is extracted rather than handed over
    whole.
    """

    def test_a_word_document_arrives_as_text(self):
        data = _docx_bytes(lambda d: d.add_paragraph("Rocks and Soils"))
        blocks = blocks_for_upload("plan.docx", data)
        assert blocks[0]["type"] == "text"
        assert "Rocks and Soils" in blocks[0]["text"]

    def test_a_table_row_stays_on_one_line(self):
        """Medium-term plans are tables. A cell torn from its row loses which
        lesson it belonged to."""

        def build(d):
            table = d.add_table(rows=1, cols=3)
            cells = table.rows[0].cells
            cells[0].text = "Lesson 1"
            cells[1].text = "types of rock"
            cells[2].text = "sorting activity"

        text = blocks_for_upload("plan.docx", _docx_bytes(build))[0]["text"]
        assert "Lesson 1 | types of rock | sorting activity" in text

    def test_document_order_survives(self):
        """The order units are taught is one of the things she is accountable
        for, so paragraphs and tables must not be flattened into two heaps."""

        def build(d):
            d.add_paragraph("Autumn 1")
            table = d.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "types of rock"
            d.add_paragraph("Assessment quiz")

        text = blocks_for_upload("plan.docx", _docx_bytes(build))[0]["text"]
        assert text.index("Autumn 1") < text.index("types of rock")
        assert text.index("types of rock") < text.index("Assessment quiz")

    def test_a_merged_cell_is_not_repeated_once_per_column(self):
        def build(d):
            table = d.add_table(rows=1, cols=3)
            row = table.rows[0]
            row.cells[0].merge(row.cells[1]).text = "Autumn term"
            row.cells[2].text = "Rocks and Soils"

        text = blocks_for_upload("plan.docx", _docx_bytes(build))[0]["text"]
        assert text.count("Autumn term") == 1

    def test_a_word_document_with_no_readable_text_is_refused(self):
        """An empty extraction reads downstream as a plan that covers nothing,
        which is the one failure this module exists to avoid."""
        with pytest.raises(UnreadableUploadError, match="plan.docx"):
            blocks_for_upload("plan.docx", _docx_bytes(lambda d: None))

    def test_a_file_that_is_not_really_a_word_document_is_refused_by_name(self):
        """An old .doc renamed to .docx is the realistic case."""
        with pytest.raises(UnreadableUploadError, match="plan.docx"):
            blocks_for_upload("plan.docx", b"not a zip archive at all")


class TestExtractionSystemPrompt:
    def test_it_forbids_adding_anything(self):
        low = EXTRACTION_SYSTEM_PROMPT.lower()
        assert "never add" in low or "do not add" in low

    def test_it_asks_for_json_only(self):
        assert "json" in EXTRACTION_SYSTEM_PROMPT.lower()


@pytest.fixture
def sent(monkeypatch):
    """Replace the model call and keep what it was given."""
    record = {"payload": dict(GOOD)}

    def fake_generate(content, system_prompt, **kwargs):
        record["content"] = content
        record["system_prompt"] = system_prompt
        record["kwargs"] = kwargs
        error = record.get("raises")
        if error is not None:
            raise error
        return record["payload"]

    monkeypatch.setattr(scheme_intake, "generate_structured_content", fake_generate)
    return record


def _read(**kwargs):
    return read_scheme_plan(
        scheme="Boost", subject="Science", year_group="Year 3", **kwargs
    )


def _texts(content):
    return [b["text"] for b in content if b["type"] == "text"]


class TestReadingTheWholePlan:
    def test_pasted_text_becomes_a_checked_plan(self, sent):
        reading = _read(pasted_text="Rocks and Soils. Covers: types of rock.")
        assert isinstance(reading.plan, SchemePlan)
        assert reading.plan.unit_title == "Rocks and Soils"
        assert reading.plan.coverage == GOOD["coverage"]

    def test_what_she_pasted_is_what_is_sent(self, sent):
        _read(pasted_text="Boost Y3 Autumn 1")
        assert any("Boost Y3 Autumn 1" in t for t in _texts(sent["content"]))

    def test_the_instruction_goes_last_after_the_source_material(self, sent):
        """Documents and images ahead of the prompt is the documented ordering,
        and the model attends to them better that way."""
        _read(
            pasted_text="Boost Y3 Autumn 1",
            uploads=[("plan.pdf", b"%PDF-1.4 fake")],
        )
        content = sent["content"]
        instruction = build_extraction_prompt(
            scheme="Boost", subject="Science", year_group="Year 3"
        )
        assert content[-1] == {"type": "text", "text": instruction}
        assert content[0]["type"] == "document"

    def test_several_photographs_are_all_sent(self, sent):
        """A printed medium-term plan is two sides of A4 more often than one."""
        _read(uploads=[("p1.jpg", b"\xff\xd8\xff a"), ("p2.jpg", b"\xff\xd8\xff b")])
        images = [b for b in sent["content"] if b["type"] == "image"]
        assert len(images) == 2

    def test_the_fidelity_instruction_is_the_system_prompt(self, sent):
        _read(pasted_text="Rocks")
        assert sent["system_prompt"] == EXTRACTION_SYSTEM_PROMPT

    def test_nothing_to_read_is_refused_before_any_call_is_made(self, sent):
        with pytest.raises(SchemePlanError, match="Nothing to read"):
            _read(pasted_text="   ")
        assert "content" not in sent, "it called Anthropic with an empty request"

    def test_a_plan_with_no_coverage_is_rejected_not_returned_empty(self, sent):
        sent["payload"] = {**GOOD, "coverage": []}
        with pytest.raises(SchemePlanError):
            _read(pasted_text="Rocks")

    def test_the_source_is_named_so_she_can_see_what_was_read(self, sent):
        reading = _read(pasted_text="Rocks", uploads=[("Y3 plan.pdf", b"%PDF")])
        assert "Y3 plan.pdf" in reading.source
        assert "paste" in reading.source.lower()


class TestWhatComesBackFlagged:
    def test_coverage_that_names_no_action_is_flagged(self, sent):
        sent["payload"] = {**GOOD, "coverage": ["properties"]}
        reading = _read(pasted_text="Rocks")
        assert [item for item, _ in reading.flagged] == ["properties"]

    def test_a_teachable_objective_is_left_alone(self, sent):
        sent["payload"] = {
            **GOOD,
            "coverage": ["compare and group rocks by their physical properties"],
        }
        assert _read(pasted_text="Rocks").flagged == []

    def test_the_extraction_may_flag_a_line_the_word_check_missed(self, sent):
        """The model reads the page; the word check only reads the sentence."""
        sent["payload"] = {
            **GOOD,
            "coverage": ["use the correct terms"],
            "vague": ["use the correct terms"],
        }
        flagged = _read(pasted_text="Rocks").flagged
        assert [item for item, _ in flagged] == ["use the correct terms"]

    def test_it_cannot_flag_something_that_is_not_in_the_coverage(self, sent):
        """A flag against a line she cannot see on screen is unanswerable, and
        an invented one would look identical to a real one."""
        sent["payload"] = {
            **GOOD,
            "coverage": ["compare and group rocks by their physical properties"],
            "vague": ["children know about soil"],
        }
        assert _read(pasted_text="Rocks").flagged == []


class TestALineMovedOutOfTheCoverage:
    """The failure found against the live API on 2026-09-02.

    Asked to flag a vague line, Claude took it out of `coverage` and put it in
    `vague` instead of listing it in both. The prompt now says not to -- but a
    dropped coverage line is the one thing this module exists to prevent, so it
    is also caught after the fact. Neither silently restored (that is how an
    invented line would get into the record she is accountable for) nor
    silently discarded: shown to her, named, for her to check.
    """

    def _reading(self, sent, vague):
        sent["payload"] = {
            **GOOD,
            "coverage": ["compare and group rocks by their physical properties"],
            "vague": vague,
        }
        return _read(pasted_text="Rocks")

    def test_a_vague_line_missing_from_the_coverage_is_reported(self, sent):
        reading = self._reading(sent, ["children know that rocks have properties"])
        assert reading.dropped == ["children know that rocks have properties"]

    def test_it_is_not_quietly_put_back_into_the_coverage(self, sent):
        reading = self._reading(sent, ["children know that rocks have properties"])
        assert reading.plan.coverage == [
            "compare and group rocks by their physical properties"
        ]

    def test_nothing_is_reported_when_the_flagged_line_is_still_there(self, sent):
        sent["payload"] = {
            **GOOD,
            "coverage": ["properties"],
            "vague": ["properties"],
        }
        assert _read(pasted_text="Rocks").dropped == []

    def test_no_vague_list_at_all_reports_nothing(self, sent):
        assert _read(pasted_text="Rocks").dropped == []

    def test_nothing_is_flagged_twice(self, sent):
        sent["payload"] = {**GOOD, "coverage": ["properties"], "vague": ["properties"]}
        assert len(_read(pasted_text="Rocks").flagged) == 1

    def test_every_flag_carries_its_reason(self, sent):
        sent["payload"] = {**GOOD, "coverage": ["properties"], "vague": ["properties"]}
        assert all(reason.strip() for _, reason in _read(pasted_text="Rocks").flagged)

    def test_the_coverage_itself_is_never_rewritten(self, sent):
        """Flag, never fix. The rewrite is a judgement she has to see."""
        sent["payload"] = {**GOOD, "coverage": ["properties"]}
        assert _read(pasted_text="Rocks").plan.coverage == ["properties"]
