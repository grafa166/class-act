"""
Prototype only — not part of the app.

Builds a single Word file that puts the CURRENT worksheet approach next to the
PROPOSED one, using identical content, so the difference can be felt by editing
rather than argued about.

Current  = single-cell tables drawing boxes + hard-coded formatting everywhere.
Proposed = ordinary paragraphs with borders/shading + real Word styles.

They look nearly the same on screen. They behave completely differently.

    cd "Class Act" && .venv/bin/python prototype/make_editable_demo.py
"""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLUE = RGBColor(0x14, 0x44, 0x8C)
BLUE_HEX = "14448C"
BLUE_FILL = "F2F6FC"
GREY_HEX = "C9D2DC"
FONT = "Arial"

OUT = "prototype/Worksheet - editing comparison.docx"


# ── low-level XML helpers ────────────────────────────────────────────────────

def _border(edge, colour, sz=8, space=6):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), colour)
    return el


def para_box(p, colour=BLUE_HEX, fill=None, sz=8):
    """Draw a box around a PARAGRAPH.

    This is the whole proposal in one function. Consecutive paragraphs carrying
    identical borders are merged by Word into a single visual box, so a box can
    hold several paragraphs and still be ordinary, typeable, reflowing text.
    """
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        bdr.append(_border(edge, colour, sz))
    pPr.append(bdr)
    if fill:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)


def answer_line(doc, style="CA Answer line"):
    """A writing line that is a real paragraph with a bottom border.

    A child writes on it; the teacher can delete it, add another, or type over
    it, and the text around it reflows. A table cell does none of that.
    """
    p = doc.add_paragraph("", style=style)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bdr.append(_border("bottom", "808080", sz=6))
    pPr.append(bdr)
    return p


def cell_shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(
        OxmlElement("w:shd")
    )
    shd = cell._tc.tcPr.find(qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def cell_borders(cell, colour):
    tcPr = cell._tc.get_or_add_tcPr()
    bdrs = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        bdrs.append(_border(edge, colour, sz=12, space=0))
    tcPr.append(bdrs)


# ── styles: the half of the fix that is invisible until she needs it ─────────

def build_styles(doc):
    """Define named styles once.

    Why it matters to her:
      * Select-all and change the font becomes one click, not a fight.
      * The Styles pane can restyle every objective box in the document at once.
      * Pasting into another document can adopt that document's look, instead of
        dragging Class Act's formatting in with it.
    """
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)

    def new(name, size, bold=False, colour=None, after=6, before=0, base="Normal"):
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles[base]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        if colour is not None:
            st.font.color.rgb = colour
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.space_before = Pt(before)
        return st

    new("CA Title", 16, True, BLUE, after=2)
    new("CA Subtitle", 11, False, RGBColor(0x4A, 0x55, 0x60), after=12)
    new("CA Section", 12, True, BLUE, after=4, before=12)
    new("CA Objective", 12, True, None, after=4)
    new("CA Body", 12, False, None, after=6)
    new("CA Answer line", 12, False, None, after=8, before=8)

    crit = new("CA Criterion", 12, False, None, after=3)
    crit.paragraph_format.left_indent = Cm(0.6)

    tag = new("CA Tag", 9, True, BLUE, after=2)
    tag.font.all_caps = True


def run(p, text, bold=False, colour=None, size=None, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.bold = bold
    r.italic = italic
    if colour is not None:
        r.font.color.rgb = colour
    if size is not None:
        r.font.size = Pt(size)
    return r


# ── the two versions of the same worksheet ───────────────────────────────────

TITLE = "Which rock for the job?"
OBJECTIVE = "We are learning to: compare two rocks and explain which is better suited to a job."
QUESTIONS = [
    ("Evidences criterion 1", "Describe each rock. Use two property words.", 2),
    ("Evidences criterion 2", "Test them. Tick the harder one, and the one that lets water through.", 1),
    ("Evidences criterion 3", "I would choose rock ____ for the pitch because …", 2),
]


def version_current(doc):
    """How it is built today: every box is a one-cell table, no styles used."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "Name ________________     Date __________", size=10)

    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    cell_borders(c, BLUE_HEX)
    cp = c.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(cp, TITLE, bold=True, colour=BLUE, size=16)

    doc.add_paragraph()

    t2 = doc.add_table(rows=1, cols=1)
    c2 = t2.cell(0, 0)
    cell_borders(c2, BLUE_HEX)
    cell_shade(c2, BLUE_FILL)
    run(c2.paragraphs[0], OBJECTIVE, size=12)

    doc.add_paragraph()

    for tag, q, lines in QUESTIONS:
        tq = doc.add_table(rows=1, cols=1)
        cq = tq.cell(0, 0)
        cell_borders(cq, GREY_HEX)
        ptag = cq.paragraphs[0]
        run(ptag, tag.upper(), bold=True, colour=BLUE, size=9)
        pq = cq.add_paragraph()
        run(pq, q, size=12)
        for _ in range(lines):
            pl = cq.add_paragraph()
            run(pl, "_" * 62, size=12)
        doc.add_paragraph()


def version_proposed(doc):
    """Same look. Ordinary paragraphs and named styles — so it can be edited."""
    p = doc.add_paragraph(style="CA Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Name ________________     Date __________")

    t = doc.add_paragraph(TITLE, style="CA Title")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_box(t, BLUE_HEX, sz=12)

    o = doc.add_paragraph(OBJECTIVE, style="CA Objective")
    para_box(o, BLUE_HEX, fill=BLUE_FILL)

    for tag, q, lines in QUESTIONS:
        # Consecutive paragraphs with the same border merge into one box in Word,
        # so this reads as a single boxed question but stays editable text.
        pt = doc.add_paragraph(tag, style="CA Tag")
        para_box(pt, GREY_HEX)
        pq = doc.add_paragraph(q, style="CA Body")
        para_box(pq, GREY_HEX)
        for _ in range(lines):
            answer_line(doc)


# ── assemble ─────────────────────────────────────────────────────────────────

def divider(doc, label, blurb):
    doc.add_page_break()
    h = doc.add_paragraph(label, style="CA Section")
    h.paragraph_format.space_before = Pt(0)
    b = doc.add_paragraph(blurb, style="CA Body")
    b.runs[0].italic = True


def main():
    doc = Document()
    build_styles(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.0)

    doc.add_paragraph("Two ways of building the same worksheet", style="CA Title")
    doc.add_paragraph(
        "They look almost identical. Try editing both and the difference is immediate.",
        style="CA Subtitle",
    )

    doc.add_paragraph("Try these four things on each version", style="CA Section")
    for line in [
        "Click in the middle of a question and press Enter to add a line.",
        "Select the whole page and change the font to something else.",
        "Copy a question and paste it into a blank Word document.",
        "Delete a writing line, then add two more.",
    ]:
        doc.add_paragraph(line, style="CA Criterion").add_run("")
    doc.paragraphs[-1]  # keep reference tidy

    divider(
        doc,
        "VERSION A — how Class Act builds it today",
        "Every box here is a one-cell table, and every word is formatted "
        "individually. There are 15 of these table-boxes in the shared worksheet "
        "code alone.",
    )
    version_current(doc)

    divider(
        doc,
        "VERSION B — proposed",
        "Same look. No tables at all: these are ordinary paragraphs with a border "
        "drawn round them, using named Word styles. Everything on this page is "
        "normal, typeable, reflowing text.",
    )
    version_proposed(doc)

    doc.add_page_break()
    doc.add_paragraph("What changes for her", style="CA Section")
    for line in [
        "Pressing Enter inside a box adds a line, instead of fighting the table.",
        "Text reflows normally — a long answer no longer stretches a cell.",
        "Select-all and change the font works, because the styles carry the formatting.",
        "Pasting into another document can adopt that document's styles instead of "
        "dragging Class Act's formatting across.",
        "Real tables stay only where content is genuinely tabular — a comparison "
        "grid, a matching activity. Those should be tables.",
    ]:
        doc.add_paragraph(line, style="CA Criterion")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
